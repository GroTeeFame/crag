from typing import List, Dict, Optional

# from chunking_evaluation.chunking import RecursiveTokenChunker
# from chunking_evaluation.utils import openai_token_count

from recursive_token_chunker import RecursiveTokenChunker, openai_token_count

import os
import re
import json
import uuid
import zipfile
import tiktoken
import pypandoc
from pprint import pprint

from docx import Document
from docx.oxml import CT_P, CT_Tbl

from colorama import Fore
from langchain.schema import Document as LangchainDocument

from config import Config

from docx_to_md import run_logic as convert_docx_to_markdown


CHUNK1 = 10
CHUNK2 = 11

CHUNK_SIZE = Config.CHUNK_SIZE
CHUNK_OVERLAP = Config.CHUNK_OVERLAP

# encoding_model = "cl100k_base"
# encoding_model = "o200k_base"

ENCODING_MODEL = "o200k_base"

HEADERS_MUCH_BUFFER = []

# Debug gating via Config.DEBUG
DEBUG = getattr(Config, "DEBUG", False)

def dprint(msg: str):
    if DEBUG:
        print(msg)


def _sha1(text: str) -> str:
    dprint(f"{'='*33} def _sha1(text) {'='*33}")
    import hashlib
    return hashlib.sha1(text.encode("utf-8")).hexdigest()

def _sha256_file(path: str) -> Optional[str]:
    try:
        import hashlib
        h = hashlib.sha256()
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(1024*1024), b''):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return None

_heading_re = re.compile(r'^\s{0,3}(#{1,6})\s+(.*)$', re.MULTILINE)

def _stable_find_positions(big: str, pieces: List[str]) -> List[tuple]:
    dprint(f"{'='*33} def _stable_find_positions(...) {'='*33}")

    """
    Find non-overlapping positions of each chunk in order.
    Prefers matches at/after the previous end; falls back to global search.
    Returns list of (start, end); (-1, -1) if not found.
    """
    positions = []
    cursor = 0
    for chunk in pieces:
        idx = big.find(chunk, cursor)
        if idx == -1:
            idx = big.find(chunk)
        if idx == -1:
            positions.append((-1, -1))
        else:
            positions.append((idx, idx + len(chunk)))
            cursor = idx + len(chunk)
    return positions

def _extract_heading_path_for_chunks(chunks: List[str]) -> List[Dict]:
    dprint(f"{'='*33} def _extract_heading_path_for_chunks(...) {'='*33}")

    """
    For each chunk, compute:
    - headings_in_chunk: all md headings found in the chunk (top->bottom).
    - heading_path: inherited H1..H6 path snapshot active at this chunk.
    """
    path = [""] * 6  # last seen titles for levels 1..6
    out: List[Dict] = []
    for chunk in chunks:
        found = []
        for m in _heading_re.finditer(chunk):
            level = len(m.group(1))
            title = m.group(2).strip()
            found.append(title)
            path[level - 1] = title
            # clear deeper levels
            for i in range(level, 6):
                if i != (level - 1):
                    path[i] = ""
        effective_path = [p for p in path if p]
        out.append({"headings_in_chunk": found, "heading_path": effective_path})
    return out

# --- Attachment/parent linkage helper ---
def _infer_attachment_metadata_from_md_path(file_path: str) -> Dict:
    dprint(f"{'='*33} def _infer_attachment_metadata_from_md_path(file_path) {'='*33}")

    """Infer linkage and department metadata from MD file path.

    Path-safe grouping:
    - attachment_group_id = sha1(relative_docx_path_without_extension)
      where relative_docx_path is under Config.DOCS_PATH.
    - Adds md_relpath (relative to Config.DOCS_PATH_MD) and docx_relpath (relative to Config.DOCS_PATH).
    - Recognizes `<parent>_dodatok_<i>.md` attachments and groups them to the same parent docx.
    """
    base = os.path.basename(file_path)
    if not base.lower().endswith(".md"):
        return {}
    stem = base[:-3]

    # Compute department and relative md path
    department = None
    md_relpath = None
    try:
        md_root = Config.DOCS_PATH_MD
        rel = os.path.relpath(file_path, md_root)
        md_relpath = rel
        if not rel.startswith(".."):
            parts = rel.split(os.sep)
            if len(parts) > 1:
                department = parts[0]
    except Exception:
        pass

    # Map MD relative path to DOCX relative path (mirrored structure)
    # For main md `<dir>/<name>.md` -> `<dir>/<name>.docx`
    # For attachments `<name>_dodatok_i.md` -> parent `<name>.docx`
    dir_rel = os.path.dirname(md_relpath) if md_relpath not in (None, "", ".") else ""
    is_attachment = False
    parent_stem = stem
    if "_dodatok_" in stem:
        parent_stem = stem.split("_dodatok_", 1)[0]
        is_attachment = True

    # Build docx_relpath relative to DOCS_PATH (mirror of md_relpath directory + parent stem)
    if dir_rel in (None, "", "."):
        docx_relpath = f"{parent_stem}.docx"
        parent_md_basename = f"{parent_stem}.md"
    else:
        docx_relpath = os.path.join(dir_rel, f"{parent_stem}.docx")
        parent_md_basename = os.path.join(dir_rel, f"{parent_stem}.md")

    # Compute content hash of the parent DOCX to use as stable, path-agnostic group id
    docx_abs = os.path.join(Config.DOCS_PATH, docx_relpath)
    content_hash = _sha256_file(docx_abs)
    if not content_hash:
        # Fallback to path-based hash if file not found (should be rare)
        norm_docx_rel_no_ext = os.path.splitext(docx_relpath.replace("\\", "/"))[0]
        content_hash = _sha1(norm_docx_rel_no_ext)
    group_id = content_hash

    out = {
        "attachment_group_id": group_id,
        "parent_basename": f"{parent_stem}.md",
        "parent_docx_basename": f"{parent_stem}.docx",
        "docx_relpath": docx_relpath.replace("\\", "/"),
        "doc_content_sha256": content_hash,
        "attachment_index": None,
    }
    if is_attachment:
        # Try to extract an index integer
        try:
            idx_str = stem.split("_dodatok_", 1)[1]
            out["attachment_index"] = int(idx_str)
        except Exception:
            out["attachment_index"] = None
        out["attachment_of"] = docx_relpath.replace("\\", "/")
    else:
        out["attachment_of"] = None

    if department is not None:
        out["department"] = department
    if md_relpath is not None:
        out["md_relpath"] = md_relpath.replace("\\", "/")
    return out

    if "_dodatok_" in stem:
        parent, idx = stem.rsplit("_dodatok_", 1)
        try:
            idx_int = int(idx)
        except Exception:
            idx_int = None
        return {
            "source_type": "embedded",
            "attachment_of": f"{parent}.md",
            **_common_fields(parent, idx_int),
        }
    else:
        # main md exported from `<parent>.docx` as `<parent>.md`
        parent = stem
        return {
            "source_type": "main",
            "attachment_of": None,
            **_common_fields(parent, None),
        }



def count_tokens2(text):
    dprint(f"{'='*33} def count_tokens2(text) {'='*33}")

    """Count tokens in a text string using tiktoken (console demo)."""
    model35="cl100k_base"
    encoder35 = tiktoken.get_encoding(model35)
    model4o="o200k_base"
    encoder4o = tiktoken.get_encoding(model4o)

    message = f"""
    Number of tokens (GPT-3.5 tokenizer): {len(encoder35.encode(text))}
    Number of tokens (GPT-4o tokenizer): {len(encoder4o.encode(text))}
    """
    return print(message)

def analyze_chunks(chunks, use_tokens=False):
    dprint(f"{'='*33} def analyze_chunks(chunks, use_tokens={use_tokens}) {'='*33}")

    if len(chunks) < 2:
        return
    CHUNK1 = 0
    CHUNK2 = 0
    # Print the chunks of interest
    print("\nNumber of Chunks:", len(chunks))
    print("\n", "="*50, f"{CHUNK1}th Chunk", "="*50,"\n", chunks[CHUNK1])
    if use_tokens:
        count_tokens2(chunks[CHUNK1])
    print("\n", "="*50, f"{CHUNK2}st Chunk", "="*50,"\n", chunks[CHUNK2])
    if use_tokens:
        count_tokens2(chunks[CHUNK2])
    # chunk1, chunk2 = chunks[199], chunks[200]
    chunk1, chunk2 = chunks[CHUNK1], chunks[CHUNK2]
    
    if use_tokens:
        encoding = tiktoken.get_encoding("cl100k_base")
        tokens1 = encoding.encode(chunk1)
        tokens2 = encoding.encode(chunk2)
        

        # Find overlapping tokens
        for i in range(len(tokens1), 0, -1):
            if tokens1[-i:] == tokens2[:i]:
                overlap = encoding.decode(tokens1[-i:])
                print("\n", "="*50, f"\nOverlapping text ({i} tokens):", overlap)
                return
        print("\nNo token overlap found")
    else:
        # Find overlapping characters
        for i in range(min(len(chunk1), len(chunk2)), 0, -1):
            if chunk1[-i:] == chunk2[:i]:
                print("\n", "="*50, f"\nOverlapping text ({i} chars):", chunk1[-i:])
                return
        print("\nNo character overlap found")



# def extract_paragraphs_with_offsets(docx_path):
#     print(f"{'='*33}  {'='*33}")

#     doc = Document(docx_path)
#     paragraphs = []
#     current_heading = None
#     offset = 0
#     para_index = 0
#     table_index = 0

#     for element in doc.element.body:
#         if isinstance(element, CT_P):
#             par = element.xpath('.//w:t')
#             paragraph_text = ''.join([node.text for node in par if node.text])
#             if paragraph_text:
#                 text = paragraph_text.strip()
#                 start = offset
#                 offset += len(text) + 1  # +1 for newline
#                 if len(text) < 80:
#                     current_heading = text
#                 metadata = {
#                     "type": "paragraph",
#                     "heading": current_heading,
#                     "index": para_index,
#                     "id": str(uuid.uuid4())
#                 }
#                 paragraphs.append({
#                     "text": text,
#                     "start": start,
#                     "end": offset,
#                     "metadata": metadata
#                 })
#                 para_index += 1

#         elif isinstance(element, CT_Tbl):
#             for row in element.xpath('.//w:tr'):
#                 row_data = []
#                 for cell in row.xpath('.//w:tc'):
#                     cell_texts = cell.xpath('.//w:t')
#                     cell_text = ''.join([node.text for node in cell_texts if node.text])
#                     row_data.append(cell_text.strip())
#                 text = ' | '.join(row_data)
#                 if text:
#                     start = offset
#                     offset += len(text) + 1
#                     metadata = {
#                         "type": "table",
#                         "heading": current_heading,
#                         "table_index": table_index,
#                         "id": str(uuid.uuid4())
#                     }
#                     paragraphs.append({
#                         "text": text,
#                         "start": start,
#                         "end": offset,
#                         "metadata": metadata
#                     })
#             table_index += 1

#     return paragraphs

def read_docx_in_order(docx_path):
    dprint(f"{'='*33} def read_docx_in_order(docx_path) {'='*33}")
    # Load the document
    doc = Document(docx_path)

    # Iterate through each element in the document body
    full_text = []
    
    for element in doc.element.body:
        if isinstance(element, CT_P):
            # Extract paragraph text
            par = element.xpath('.//w:t')
            paragraph_text = ''.join([node.text for node in par if node.text])
            if paragraph_text:
                full_text.append(paragraph_text.strip())

        elif isinstance(element, CT_Tbl):
            # Extract table data
            table = element
            for row in table.xpath('.//w:tr'):
                row_data = []
                for cell in row.xpath('.//w:tc'):
                    cell_texts = cell.xpath('.//w:t')
                    cell_text = ''.join([node.text for node in cell_texts if node.text])
                    row_data.append(cell_text.strip())
                full_text.append(' | '.join(row_data))

    # Join all collected text into a single string
    complete_text = '\n'.join(full_text)
    
    return complete_text



def extract_embedded_files(docx_path, output_dir):
    dprint(f"{'='*33} def extract_embedded_files(docx_path, output_dir) {'='*33}")

    extracted_files = []
    # Open the .docx file as a zip archive
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        # Look for embedded files in the 'word/embeddings' directory
        embeddings_dir = 'word/embeddings/'
        # Ensure the output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Extract each embedded file to the output directory
        for file in docx_zip.namelist():
            if file.startswith(embeddings_dir):
                docx_zip.extract(file, output_dir)
                extracted_files.append(file)
                dprint(f"Extracted: {file}")

    return extracted_files

# def read_full_document(file_path):
#     print(f"{'='*33}  {'='*33}")

#     print("inside read_full_document()")

#     output_dir = 'extracted_files/'
#     embeddings_dir = 'word/embeddings/'
#     full_embed_path = f"{output_dir}/{embeddings_dir}"
#     # extracted_files = extract_embedded_files(file_path, "extracted_files")
#     extracted_files = extract_embedded_files(file_path, output_dir)
#     # document_text = convert_docx_to_markdown(file_path)
#     document_text = read_docx_in_order(file_path)
#     if extracted_files:
#         for i, extracted_file_path in enumerate(extracted_files, start=1):
#             print('read_full_document() - for i, extracted_file_path in enumerate(extracted_files): ')
#             print('\n\n')
#             print(f"extracted_file_path : {extracted_file_path}")
#             print('\n\n')
#             if extracted_file_path.endswith(".docx"):
#                 embedded_path = os.path.join(output_dir, extracted_file_path)
#                 appended = read_docx_in_order(embedded_path)
#                 document_text = f"{document_text} \n# Додаток №{i}: \n {appended}"
#             # if extracted_file_path.endswith(".docx"):
#             #     document_text = f"{document_text} \n# Додаток №{i}: \n {read_docx_in_order(f"{output_dir}{extracted_file_path}")}"
#                 # document_text = f"{document_text} \n# Додаток №{i}: \n {convert_docx_to_markdown(f"{output_dir}{extracted_file_path}")}"

#         # for filename in os.listdir(f"{output_dir}/{embeddings_dir}"):
#         for filename in os.listdir(full_embed_path):
#             file_path = os.path.join(full_embed_path, filename)
#             if os.path.isfile(file_path):
#                 os.remove(file_path)
#                 print(filename, "is removed")
    
#     return document_text


def read_full_document_md(file_path):
    dprint(f"{'='*33} def read_full_document_md(file_path) {'='*33}")

    output_dir = 'extracted_files/'
    embeddings_dir = 'word/embeddings/'
    full_embed_path = f"{output_dir}/{embeddings_dir}"
    # extracted_files = extract_embedded_files(file_path, "extracted_files")
    extracted_files = extract_embedded_files(file_path, output_dir)
    document_text = convert_docx_to_markdown(file_path)
    # document_text = read_docx_in_order(file_path)
    if extracted_files:
        for i, extracted_file_path in enumerate(extracted_files, start=1):
            dprint('read_full_document() - iterating embedded files')
            dprint(f"extracted_file_path : {extracted_file_path}")
            if extracted_file_path.endswith(".docx"):
                embedded_path = os.path.join(output_dir, extracted_file_path)
                appended = convert_docx_to_markdown(embedded_path)
                document_text = f"{document_text} \n# Додаток №{i}: \n {appended}"
            # if extracted_file_path.endswith(".docx"):
            #     # document_text = f"{document_text} \n# Додаток №{i}: \n {read_docx_in_order(f"{output_dir}{extracted_file_path}")}"
            #     document_text = f"{document_text} \n# Додаток №{i}: \n {convert_docx_to_markdown(f"{output_dir}{extracted_file_path}")}"

        # for filename in os.listdir(f"{output_dir}/{embeddings_dir}"):
        for filename in os.listdir(full_embed_path):
            file_path = os.path.join(full_embed_path, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
                dprint(f"Removed embedded file: {filename}")
    
    return document_text


def recursive_tokens_chunking(document):
    dprint(f"{'='*33} def recursive_tokens_chunking(document) {'='*33}")

    recursive_token_chunker = RecursiveTokenChunker(
        chunk_size=CHUNK_SIZE,              # token length
        chunk_overlap=CHUNK_OVERLAP,        # token overlap
        length_function=openai_token_count,
        separators=["\n###### ", "\n##### ", "\n#### ", "\n### ", "\n## ", "\n# ", "\n\n", ".", "\n", " ", ""]
    )

    recursive_token_chunks = recursive_token_chunker.split_text(document)

    if DEBUG:
        analyze_chunks(recursive_token_chunks, use_tokens=True)
    return recursive_token_chunks


def enrich_chunks_with_metadata(
    chunks,
    document_text,
    encoding_model,
    filename,
    headings_lists=None,
    heading_paths=None,
    extra_metadata: Dict = None,
):
    dprint("enrich_chunks_with_metadata()")

    enc_o200k = tiktoken.get_encoding("o200k_base")
    enc_35 = tiktoken.get_encoding("cl100k_base")
    positions = _stable_find_positions(document_text, chunks)

    # try to get file stats (best-effort)
    file_size = None
    file_mtime_iso = None
    try:
        p = filename if os.path.isabs(filename) else os.path.join(os.getcwd(), filename)
        st = os.stat(p)
        file_size = st.st_size
        from datetime import datetime
        file_mtime_iso = datetime.fromtimestamp(st.st_mtime).isoformat()
    except Exception:
        pass

    enriched = []
    for i, chunk in enumerate(chunks):
        start, end = positions[i]
        token_count_o200k = len(enc_o200k.encode(chunk))
        token_count_35 = len(enc_35.encode(chunk))
        word_count = len(chunk.split())
        headings_in_chunk = headings_lists[i] if headings_lists else []
        heading_path = heading_paths[i] if heading_paths else []
        section_title = heading_path[-1] if heading_path else (headings_in_chunk[-1] if headings_in_chunk else None)

        def _line_at(pos: int) -> int:
            if pos is None or pos < 0:
                return None
            return document_text.count("\n", 0, pos) + 1

        meta = {
            "chunk_id": str(uuid.uuid4()),
            "doc_id": _sha1(filename),
            "checksum_sha1": _sha1(chunk),

            "chunk_index": i,
            "char_range": [start, end],
            "line_span": [_line_at(start), _line_at(end)],

            "token_count_o200k": token_count_o200k,
            "token_count_cl100k": token_count_35,
            "word_count": word_count,

            "headings": headings_in_chunk,
            "heading_path": heading_path,
            "section_title": section_title,

            "filename": filename,
            "file_size": file_size,
            "file_mtime_iso": file_mtime_iso,
            "encoding_model": encoding_model,
        }

        # serialize list fields for vectorstores that reject non-scalars
        meta["char_range_str"] = json.dumps(meta["char_range"], ensure_ascii=False)
        meta["line_span_str"] = json.dumps(meta["line_span"], ensure_ascii=False)
        
        if extra_metadata:
            # do not overwrite the core keys unless explicitly provided
            for k, v in extra_metadata.items():
                meta.setdefault(k, v)

        enriched.append({
            "text": chunk,
            "metadata": meta
        })
    return enriched



def convert_docx_to_markdown(docx_path: str) -> str:
    print(f"{'='*33} def convert_docx_to_markdown(docx_path: str) -> str: {'='*33}")

    return pypandoc.convert_file(docx_path, 'markdown', format='docx')

# Wrap main logic in function

# def parse_docx_to_chunks_md(file_path: str) -> List[Dict]:
#     print(f"{'='*33}  {'='*33}")

#     # paragraphs = extract_paragraphs_with_offsets(file_path)
#     # document_text = read_docx_in_order(file_path)
#     document_text = convert_docx_to_markdown(file_path)
#     chunks = recursive_tokens_chunking(document_text)
#     enriched_chunks = enrich_chunks_with_metadata(
#         chunks,
#         document_text,
#         # paragraphs,
#         ENCODING_MODEL,
#         os.path.basename(file_path)
#     )
#     return enriched_chunks




# def extract_section_heading(chunk: str) -> str:
# def extract_section_heading(chunk: str, headers_much_buffer):
#     print(f"{'='*33}  {'='*33}")

#     """
#     Extracts the most likely section heading from a chunk of text.
#     Looks for the last Markdown-style or numbered section title.
#     """
#     lines = chunk.strip().split('\n')
#     headings = []
#     md_headings = []
#     # Patterns
#     md_heading_pattern = re.compile(r'^\s{0,3}#{1,6}\s+(.*)')
#     numbered_heading_pattern = re.compile(r'^\s{0,3}(\d+(\.\d+)*)([.)]|\s+)(.+)')

#     for line in lines:
#         line = line.strip()
#         md_match = md_heading_pattern.match(line)
#         num_match = numbered_heading_pattern.match(line)

#         if md_match:
#             print(f"md_match : {md_match.group(1).strip()}")
#             headings.append(md_match.group(1).strip())
#             md_headings.append(md_match.group(1).strip())

#         elif num_match:
#             heading_text = f"{num_match.group(1)} {num_match.group(4).strip()}"
#             headings.append(heading_text)

#     # Return the last heading found (most relevant for chunk)

#     if md_headings:
#         headers_much_buffer = md_headings
#     else:
#         print(Fore.RED + "+++ INSERTING HEADERS FROM PREVIOUS CHUNK +++\n\n" + Fore.RESET)
#         for head in headers_much_buffer:
#             headings.insert(0, head)

#     print(f"headings : {headings}")

#     return headings, headers_much_buffer


# def extract_headings(chunks):
#     print(f"{'='*33}  {'='*33}")

#     """
#     Returns (headings_lists, unchanged_chunks).
#     headings_lists[i] contains md headings found in chunks[i].
#     """
#     info = _extract_heading_path_for_chunks(chunks)
#     headings_lists = [d["headings_in_chunk"] for d in info]
#     return headings_lists, chunks


def get_chunks_with_meta_md_only(file_path):
    dprint(f"{'='*33} def get_chunks_with_meta_md_only(file_path) {'='*33}")

    with open(file_path, "r", encoding="utf-8") as md_file:
        md_document_text = md_file.read()
    chunks = recursive_tokens_chunking(md_document_text)

    info = _extract_heading_path_for_chunks(chunks)
    headings_lists = [d["headings_in_chunk"] for d in info]
    heading_paths = [d["heading_path"] for d in info]

    extra_meta = _infer_attachment_metadata_from_md_path(file_path)

    enriched_chunks = enrich_chunks_with_metadata(
        chunks,
        md_document_text,
        ENCODING_MODEL,
        os.path.basename(file_path),
        headings_lists=headings_lists,
        heading_paths=heading_paths,
        extra_metadata=extra_meta,
    )
    return enriched_chunks


def get_chunks_from_all_documents_in_directory_md_only(documents_dir):
    dprint(f"{'='*33} def get_chunks_from_all_documents_in_directory_md_only(documents_dir) {'='*33}")
    """Recursively walk an MD root and return LangchainDocument chunks.

    Preserves metadata enrichment (including department inferred from path).
    """
    all_chunks = []
    for root, _dirs, files in os.walk(documents_dir):
        for filename in files:
            if not filename.lower().endswith(".md"):
                continue
            full_path = os.path.join(root, filename)
            enriched_chunks = get_chunks_with_meta_md_only(full_path)
            for chunk in enriched_chunks:
                all_chunks.append(LangchainDocument(
                    page_content=chunk["text"],
                    metadata=chunk["metadata"]
                ))
    return all_chunks



if __name__ == "__main__":
    print(f"{'='*33} if __name__ == __main__: parse_to_chunks.py {'='*33}")


    # file_path = 'documents/IRU2.docx'
    file_path = 'documents/_poryadok_koristuvannya_epb_v3-1_final.docx'

    enriched_chunks = get_chunks_with_meta(file_path)

    total_tokens = 0

    for i, chunk in enumerate(enriched_chunks):
        total_tokens += chunk['metadata']['token_count']
        print(Fore.BLUE + f"-----------------------------------------------------------------------------------------------enriched_chunk {i}---chunk_size: {chunk['metadata']['token_count']}" + Fore.RESET)
        # print(chunk)

        pprint(chunk)
        # for key, value in chunk.items():
        #     print(f"{key}: {value}")
        #     print("----------------")
        # for item in chunk:
        #     print(item, ":", enriched_chunks)
        # # print('++++++++++++++++++++++++++')
        # # count_tokens2(chunk)
        # # print('++++++++++++++++++++++++++')
        # for x in cars:
        #     print (x)
        #     for y in cars[x]:
        #         print (y,':',cars[x][y])
        print(Fore.BLUE + f"-----------------------------------------------------------------------------------------------end of enriched_chunk {i} ---chunk_size: {chunk['metadata']['token_count']}" + Fore.RESET)

    print(Fore.YELLOW + f"TOTAL TOKENS : {total_tokens}" + Fore.RESET)
    # with open(f"{file_path[10:-4]}.txt", "w") as f:
    #     f.write(" ".join(enriched_chunks))
    #     print(f'file was writen as : {file_path[10:-4]}')
    # print(f"PARAGRAPHS: {paragraphs}")
    # for i, paragraph in enumerate(paragraphs):
    #     print(f"------Paragraph {i}-------")
    #     print(paragraph)
    #     print('--------------------------')


    # print(f"document_text: {document_text}")
    # print(f"chunks: {chunks}")
    # for i, chunk in enumerate(chunks):
    #     print(f"-----------------------------------------------------------------------------------------------Chunk {i} ---")
    #     print(chunk)
    #     print('++++++++++++++++++++++++++')
    #     count_tokens2(chunk)
    #     print('++++++++++++++++++++++++++')
    #     meta_for_chunk, HEADERS_MUCH_BUFFER = extract_section_heading(chunk, HEADERS_MUCH_BUFFER)
    #     print(f"{Fore.BLUE}\nmeta_for_chunk : {meta_for_chunk}{Fore.RESET}")
    #     print(f"{Fore.GREEN}HEADERS_MUCH_BUFFER : {HEADERS_MUCH_BUFFER}{Fore.RESET}")
    #     # print(extract_section_heading(chunk, HEADERS_MUCH_BUFFER))
    #     print(f"------------------------------------------------------------------------------------------------end of chunk {i} -----------")

    # analyze_chunks(chunks)

    # print(f"enriched_chunks: {enriched_chunks}")
    # for i, chunk in enumerate(enriched_chunks):
    #     print(f"-----------------------------------------------------------------------------------------------enriched_chunk {i}---")
    #     print(chunk)
    #     # print('++++++++++++++++++++++++++')
    #     # count_tokens2(chunk)
    #     # print('++++++++++++++++++++++++++')
    #     print(f"------------------------------------------------------------------------------------------------end of enriched_chunk {i} -----------")
