import os
import json
import uuid
import queue
import threading
import time
import random
import unicodedata
import re
from datetime import datetime
import concurrent.futures
from typing import Dict, Any, Optional
import logging
from logging.handlers import RotatingFileHandler
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, send_file, Response, abort, session, redirect, url_for
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
try:
    from rq import Queue as RQQueue
except Exception:
    RQQueue = None
try:
    import redis as _redis
except Exception:
    _redis = None

# === reuse your existing code ===
load_dotenv()
from config import Config
from rag_logic import get_llm, get_embedding_model, should_rebuild_vectorstore, rebuild_vector_store
from read_split_nbu import split_docx_to_question_with_ids
from save_results import add_comment_to_paragraphs, sanitize_markdown
# Prefer a bundled modern SQLite (pysqlite3) if system sqlite3 is too old
try:
    __import__("pysqlite3")
    import sys as _sys
    _sys.modules["sqlite3"] = _sys.modules.pop("pysqlite3")
except Exception:
    pass

from langchain_chroma import Chroma
from chromadb import PersistentClient
from chromadb.config import Settings as ChromaSettings
from langchain.schema import Document as LangchainDocument
from langchain_community.retrievers import BM25Retriever

from rag_functions import (
    rerank_by_keyword_overlap,
    rrf_fuse, build_bm25_retriever, pack_context, cap_per_ird,
    build_packed_context, source_label_from_meta
)
from cache_manifest import cache_lookup, cache_store, sha256_file, cache_prune_missing

from openai import RateLimitError
try:
    # Broader OpenAI client exceptions (v1.x)
    from openai import APIError, APITimeoutError, APIConnectionError
except Exception:
    APIError = APITimeoutError = APIConnectionError = Exception

app = Flask(__name__)
app.config.from_object(Config)
# Session secret (for simple auth)
app.secret_key = os.getenv("SECRET_KEY", "change-me")
app.config.setdefault("SESSION_COOKIE_HTTPONLY", True)
app.config.setdefault("SESSION_COOKIE_SAMESITE", "Lax")

# Trust proxy headers (X-Forwarded-For/Proto) so limiter sees real client IP
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1)

def _setup_logging():
    """Ensure a rotating file handler is attached to the root logger.

    - JSON-ish line format to simplify ingestion
    - File logs rotate at ~10MB with 5 backups
    - Controlled via env: LOG_DIR (default: ./logs)
    - Keeps existing handlers, but guarantees the file handler exists
    """
    log_dir = os.getenv("LOG_DIR", "logs")
    os.makedirs(log_dir, exist_ok=True)
    fmt = '{"t":"%(asctime)s","level":"%(levelname)s","msg":%(message)r}'

    root = logging.getLogger()
    # Always set a sane level; do not return early if handlers already exist
    root.setLevel(logging.DEBUG if str(os.getenv("DEBUG", "0")).lower() in ("1","true","yes") else logging.INFO)

    # Rotating file handler: attach if not already present for this path
    want_path = os.path.abspath(os.path.join(log_dir, "app.log"))
    have_file = False
    for h in list(root.handlers):
        if isinstance(h, RotatingFileHandler) and getattr(h, "baseFilename", None) == want_path:
            have_file = True
            # ensure formatter
            h.setFormatter(logging.Formatter(fmt))
    if not have_file:
        fh = RotatingFileHandler(want_path, maxBytes=10*1024*1024, backupCount=5)
        fh.setLevel(logging.INFO)
        fh.setFormatter(logging.Formatter(fmt))
        root.addHandler(fh)

    # Optional console handler: add one if none exists (useful for local dev)
    if not any(isinstance(h, logging.StreamHandler) and not isinstance(h, RotatingFileHandler) for h in root.handlers):
        ch = logging.StreamHandler()
        ch.setLevel(logging.INFO)
        ch.setFormatter(logging.Formatter(fmt))
        root.addHandler(ch)

_setup_logging()

def _ensure_runtime_directories():
    """Create required folders so first-run doesn't crash on missing paths."""
    try:
        os.makedirs(getattr(Config, 'UPLOADS_PATH', 'uploads'), exist_ok=True)
        os.makedirs(getattr(Config, 'RESULT_PATH', 'docx_results'), exist_ok=True)
        os.makedirs(getattr(Config, 'DOCS_PATH', 'documents'), exist_ok=True)
        os.makedirs(getattr(Config, 'DOCS_PATH_MD', os.path.join('documents', 'documents_converted_to_md')), exist_ok=True)
        # Vectorstore parent dir
        db_dir = getattr(Config, 'DB_DIR', None)
        if db_dir:
            os.makedirs(db_dir, exist_ok=True)
        else:
            parent = os.path.dirname(getattr(Config, 'DB_NAME', 'vectorstore/db'))
            if parent:
                os.makedirs(parent, exist_ok=True)
        # Checkpoints and other working folders
        os.makedirs(getattr(Config, 'CHECKPOINT_DIR', os.path.join('work_results', 'checkpoints')), exist_ok=True)
        os.makedirs('work_results', exist_ok=True)
        os.makedirs('full_prompt', exist_ok=True)
        os.makedirs('extracted_files', exist_ok=True)
        # Optional NBU folder for local usage
        os.makedirs(getattr(Config, 'NBU_DOCS_PATH', 'NBU'), exist_ok=True)
    except Exception as e:
        logging.warning(f"Failed to ensure runtime directories: {e}")

_ensure_runtime_directories()
# Best-effort: prune stale processed-manifest entries at startup
try:
    cache_prune_missing()
except Exception:
    pass

# Rate limiting (choose storage via Config; supports Redis or memory)
limiter = Limiter(get_remote_address, app=app, storage_uri=getattr(Config, "LIMITER_STORAGE_URI", "memory://"))

# Debug gating across the app
DEBUG = getattr(Config, "DEBUG", False)

def dprint(msg: str):
    if DEBUG:
        print(msg)

def _is_docx_filename(name: str) -> bool:
    try:
        return os.path.splitext(name)[1].lower() == ".docx"
    except Exception:
        return False

def _safe_ascii_stem(text: str, fallback: str = "file") -> str:
    try:
        norm = unicodedata.normalize("NFKD", text)
        ascii_only = norm.encode("ascii", "ignore").decode("ascii")
        ascii_only = re.sub(r"[^A-Za-z0-9._-]+", "_", ascii_only).strip("._-")
        return ascii_only or fallback
    except Exception:
        return fallback

def _invoke_with_backoff(llm, prompt: str, max_retries: int = 5, base_delay: float = 2.0):
    """Invoke LLM with exponential backoff on RateLimitError.

    Args:
        llm: LangChain LLM instance with .invoke()
        prompt: Prompt string
        max_retries: Maximum retry attempts
        base_delay: Initial delay in seconds
    Returns:
        LLM response
    Raises:
        The last exception if all retries fail
    """
    attempt = 0
    while True:
        try:
            return llm.invoke(prompt)
        except (RateLimitError, APITimeoutError, APIConnectionError, APIError, TimeoutError) as e:
            attempt += 1
            if attempt > max_retries:
                raise
            # Exponential backoff with jitter
            delay = base_delay * (2 ** (attempt - 1))
            delay = min(delay, 60)
            jitter = random.uniform(0, base_delay * 0.25)
            dprint(f"Rate limit hit; retry {attempt}/{max_retries} in {round(delay + jitter, 2)}s")
            time.sleep(delay + jitter)

# ---------- Global state ----------
DB = None
LLM_CHAT = None
LLM_DOC = None

# Optional Redis connection for task state
USE_REDIS = getattr(Config, "USE_REDIS_TASKS", False)
REDIS = None
if USE_REDIS and _redis is not None:
    try:
        REDIS = _redis.from_url(getattr(Config, "REDIS_URL", "redis://127.0.0.1:6379/0"), decode_responses=True)
        # Simple ping to validate early
        REDIS.ping()
    except Exception as e:
        logging.warning(f"Redis configured but not reachable: {e}. Falling back to in-memory tasks.")
        REDIS = None
        USE_REDIS = False

# task state helpers: in-memory fallback or Redis-backed
# In-memory: task_id -> {status, total, done, started_at, finished_at, filename_in, filename_out, stream_queue}
TASKS = {}

def _task_key(task_id: str) -> str:
    return f"task:{task_id}"

def _task_events_key(task_id: str) -> str:
    return f"task_events:{task_id}"

def task_create(task_id: str, data: dict):
    if USE_REDIS and REDIS is not None:
        # Store as hash; convert values to strings where needed
        m = {k: (str(v) if isinstance(v, (int, float)) else (v or "")) for k, v in data.items() if k != "stream_queue"}
        REDIS.hset(_task_key(task_id), mapping=m)
        REDIS.expire(_task_key(task_id), 7 * 24 * 3600)  # 7 days retention
        # Optional index of tasks
        try:
            REDIS.sadd("tasks:all", task_id)
        except Exception:
            pass
    else:
        q = queue.Queue()
        data = {**data, "stream_queue": q}
        TASKS[task_id] = data

def task_get(task_id: str) -> Optional[dict]:
    if USE_REDIS and REDIS is not None:
        h = REDIS.hgetall(_task_key(task_id))
        if not h:
            return None
        # Coerce types for known numeric fields
        def as_int(x):
            try:
                return int(x)
            except Exception:
                return 0
        return {
            "status": h.get("status"),
            "total": as_int(h.get("total")),
            "done": as_int(h.get("done")),
            "started_at": float(h.get("started_at")) if h.get("started_at") else None,
            "finished_at": float(h.get("finished_at")) if h.get("finished_at") else None,
            "filename_in": h.get("filename_in"),
            "filename_out": h.get("filename_out"),
            "queue_pos": as_int(h.get("queue_pos")) if h.get("queue_pos") else None,
        }
    return TASKS.get(task_id)

def task_merge(task_id: str, patch: dict):
    if USE_REDIS and REDIS is not None:
        m = {k: (str(v) if isinstance(v, (int, float)) else (v if v is not None else "")) for k, v in patch.items()}
        if m:
            REDIS.hset(_task_key(task_id), mapping=m)
    else:
        if task_id in TASKS:
            TASKS[task_id].update(patch)

def task_push_event(task_id: str, msg: dict):
    if USE_REDIS and REDIS is not None:
        try:
            REDIS.rpush(_task_events_key(task_id), json.dumps(msg, ensure_ascii=False))
            REDIS.expire(_task_events_key(task_id), 7 * 24 * 3600)
        except Exception:
            pass
    else:
        t = TASKS.get(task_id)
        if t:
            t["stream_queue"].put(msg)

### TODO: Semaphore for too many documents <==

# ---------- Concurrency controls for doc processing ----------
# Max number of concurrent heavy doc jobs (Pandoc/IO/etc.)
DOC_MAX_CONCURRENCY = int(os.getenv("DOC_MAX_CONCURRENCY", "4"))
# Max number of queued jobs waiting to start
DOC_QUEUE_MAX       = int(os.getenv("DOC_QUEUE_MAX", "50"))

# Gate running jobs
_DOC_SEM   = threading.BoundedSemaphore(DOC_MAX_CONCURRENCY)
# Worker pool that executes doc jobs
_EXEC      = concurrent.futures.ThreadPoolExecutor(
    max_workers=DOC_MAX_CONCURRENCY,
    thread_name_prefix="docjob"
)
# Bounded queue for pending tasks (don’t flood memory)
_PENDING_Q = queue.Queue(maxsize=DOC_QUEUE_MAX)

def _run_task_guarded(task_id: str):
    dprint(f"{'='*33} def _run_task_guarded(task_id: str): {'='*33}")
    
    """
    Acquire the semaphore (limits concurrent jobs), mark status running,
    call the original heavy processor, then release the semaphore.
    """
    dprint(f"inside _run_task_guarded()  : TASK_ID: {task_id}")
    _DOC_SEM.acquire()
    try:
        # mark running
        task_merge(task_id, {"status": "running", "started_at": time.time()})
        task_push_event(task_id, {"status": "running"})
        # process
        _process_document_task(task_id)
    finally:
        _DOC_SEM.release()

def _dispatcher_loop():
    dprint(f"{'='*33} def _dispatcher_loop(): {'='*33}")
    
    """
    Pull task_ids from the bounded queue and submit them to the pool.
    Keeps pool work queue bounded and provides back‑pressure.
    """
    while True:
        task_id = _PENDING_Q.get()
        try:
            _EXEC.submit(_run_task_guarded, task_id)
        except Exception as e:
            t = TASKS.get(task_id)
            if t:
                t["status"] = "error"
                t["finished_at"] = time.time()
                t["error"] = f"submit_failed: {e}"
                t["stream_queue"].put({"status": "error", "error": t["error"]})
        finally:
            _PENDING_Q.task_done()

# Start dispatcher thread once
if not getattr(Config, "USE_RQ_TASKS", False):
    _DISPATCHER = threading.Thread(target=_dispatcher_loop, name="doc-dispatcher", daemon=True)
    _DISPATCHER.start()


### TODO: Semaphore for too many documents ==>





# ---------- Bootstrapping ----------
def init_vectorstore_and_models():
    dprint(f"{'='*33} def init_vectorstore_and_models(): {'='*33}")
    
    global DB, LLM_CHAT, LLM_DOC
    # Build or open DB
    if should_rebuild_vectorstore():
        rebuild_vector_store()
    client = PersistentClient(path=Config.DB_NAME)
    DB = Chroma(
        client=client,
        collection_name=Config.COLLECTION_NAME,
        embedding_function=get_embedding_model(),
    )
    # LLMs
    LLM_CHAT = get_llm(kind='question')
    LLM_DOC = get_llm(kind='document')

init_vectorstore_and_models()

# ---------- Small helpers ----------
def iter_docs_from_chroma(db, page_size: int = 1000):
    dprint(f"{'='*33} def iter_docs_from_chroma(db, page_size={page_size}): {'='*33}")
    """Yield LangchainDocument objects from Chroma in pages to keep memory bounded."""
    col = db._collection
    try:
        total = col.count()
        offset = 0
        while offset < total:
            res = col.get(include=["documents", "metadatas"], limit=page_size, offset=offset)
            docs = res.get("documents", []) or []
            metas = res.get("metadatas", []) or []
            got = 0
            for d, m in zip(docs, metas):
                got += 1
                yield LangchainDocument(page_content=d, metadata=(m or {}))
            if not got:
                break
            # Prefer ids length if available
            ids = res.get("ids", [])
            offset += len(ids) if ids else got
    except Exception:
        # Fallback: page by explicit IDs
        id_res = col.get(include=[])
        ids = id_res.get("ids", []) if isinstance(id_res, dict) else []
        for i in range(0, len(ids), page_size):
            sl = ids[i:i+page_size]
            res = col.get(ids=sl, include=["documents", "metadatas"])
            for d, m in zip(res.get("documents", []) or [], res.get("metadatas", []) or []):
                yield LangchainDocument(page_content=d, metadata=(m or {}))

def _make_subqueries(q: str) -> list[str]:
    # Legacy alias; use rag_functions.make_subqueries elsewhere
    from rag_functions import make_subqueries
    return make_subqueries(q)

def _color_from_answer(ans_emoji: str) -> str:
    dprint(f"{'='*33} def _color_from_answer(ans_emoji: str) -> str: {'='*33}")
    
    # map to save_results color keys
    if ans_emoji == "✅":
        return "green"
    if ans_emoji == "❌":
        return "red"
    return "yellow"

# ---------- Web ----------
@app.route("/")
def index():
    dprint(f"{'='*33} def index(): {'='*33}")
    
    return render_template("index.html",
                           model=Config.GPT_MODEL,
                           db_path=Config.DB_NAME,
                           collection=Config.COLLECTION_NAME,
                           logged_in=bool(session.get("logged_in")))

# --- Simple login for Corpus management ---
@app.route("/login", methods=["GET", "POST"])
@limiter.limit("5/minute")
def login():
    if request.method == "POST":
        u = (request.form.get("username") or "").strip()
        p = (request.form.get("password") or "").strip()
        if u == getattr(Config, 'LOGIN', '') and p == getattr(Config, 'PASSWORD', ''):
            session["logged_in"] = True
            return redirect(url_for("index"))
        return render_template("login.html", error="Невірні дані входу")
    return render_template("login.html")

@app.get("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))

@app.before_request
def _guard_corpus_api():
    # Protect corpus management APIs behind simple login
    path = request.path or ""
    if path.startswith("/api/corpus/"):
        if not session.get("logged_in"):
            # JSON for APIs
            return jsonify({"error": "auth_required"}), 401

# --- Chat API ---
@app.post("/api/chat")
@limiter.limit("30/minute")
def api_chat():
    dprint(f"{'='*33} def api_chat(): {'='*33}")
    
    req_data = request.get_json(silent=True) or {}
    user_msg = req_data.get("message", "").strip()
    if not user_msg:
        return jsonify({"error": "empty message"}), 400

    # Basic retrieval (dense + keyword rerank), pack context
    dense = DB.as_retriever(search_kwargs={"k": Config.RETURN_CHUNK_FOR_SINGLE_QUESTION})
    raw = dense.invoke(user_msg)
    ranked = rerank_by_keyword_overlap(user_msg, raw)

    # Debug dump of matched metadata (hidden unless DEBUG)
    if DEBUG:
        logging.debug("\n" + '=='*80 + "\n" + '=='*80)
        for do in ranked:
            logging.debug(f"meta={do.metadata}")
            logging.debug("-"*40)
        logging.debug('=='*80)
        logging.debug('=='*80)

    # Format context for your prompt style
    ctx = "\n\n".join(
        f"===\n[Джерело: {_source_label_from_meta(d.metadata)}]\n{d.page_content}\n===" for d in ranked
    )
    system_prompt = Config.system_prompt_question_loop
    dprint(f"system_prompt : {system_prompt}")

    prompt = f"""{system_prompt}
Подроблене запитання з контекстом:
{ctx}

Основне запитання:
{user_msg}

Відповідь українською:
"""

    try:
        res = _invoke_with_backoff(LLM_CHAT, prompt)
    except Exception as e:
        # Surface a clear error for transient/backoff failures
        return jsonify({
            "error": "llm_unavailable",
            "message": str(e)
        }), 503
    dprint(f"RES: {res}")

    out = res.content if hasattr(res, "content") else str(res)
    dprint(f"OUT: {out}")

    text_answer = out
    sources: list[Any] = []
    try:
        parsed = json.loads(out)
        text_answer = parsed.get("text", out)
        sources = parsed.get("source", [])
    except Exception:
        pass  # if LLM returns non-JSON, still continue

    dprint(f"parsed: {locals().get('parsed', None)}")
    dprint(f"text_answer: {text_answer}")
    dprint(f"sources: {sources}")

    meta = getattr(res, "response_metadata", {}) or {}
    return jsonify({
        "answer": text_answer,
        "sources": sources,
        "usage": meta.get("token_usage", {})
    })
    # return jsonify({
    #     "answer": out,
    #     "usage": meta.get("token_usage", {})
    # })

# --- Upload & long processing ---
@app.post("/api/upload")
@limiter.limit("5/minute")
def api_upload():
    dprint(f"{'='*33} def api_upload(): {'='*33}")
    
    if "file" not in request.files:
        return jsonify({"error": "no file"}), 400
    f = request.files["file"]
    # Validate extension against allowlist
    if not f or not f.filename:
        return jsonify({"error": "no filename"}), 400
    orig_name = f.filename or ""
    base_name = os.path.basename(orig_name)
    if not _is_docx_filename(base_name):
        return jsonify({"error": "only .docx is supported"}), 400

    os.makedirs(Config.UPLOADS_PATH, exist_ok=True)
    os.makedirs(Config.RESULT_PATH, exist_ok=True)

    # unique in/out paths
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    in_path = os.path.join(Config.UPLOADS_PATH, f"{ts}_{base_name}")
    stem = os.path.splitext(base_name)[0]
    out_path = os.path.join(Config.RESULT_PATH, f"AI_{ts}_{stem}.docx")
    f.save(in_path)

    # Deduplicate identical uploads (same content) to avoid double processing
    def _sha256(p: str) -> str:
        import hashlib
        h = hashlib.sha256()
        with open(p, 'rb') as _fp:
            for chunk in iter(lambda: _fp.read(1024 * 1024), b''):
                h.update(chunk)
        return h.hexdigest()

    file_hash = _sha256(in_path)

    # Persistent cache: if an identical file was processed before, return a completed task immediately
    cached_out = None
    try:
        cached_out = cache_lookup(file_hash)
    except Exception:
        cached_out = None
    if cached_out and os.path.exists(cached_out):
        task_id = str(uuid.uuid4())
        # create an immediate-done task so UI can attach SSE and show 'done'
        task_create(task_id, {
            "status": "done",
            "total": 1,
            "done": 1,
            "started_at": time.time(),
            "finished_at": time.time(),
            "filename_in": in_path,
            "filename_out": cached_out,
        })
        # push done event for SSE consumers
        task_push_event(task_id, {"status": "done", **_task_snapshot(task_id)})
        return jsonify({"task_id": task_id, "cache_hit": True})
    if USE_REDIS and REDIS is not None:
        try:
            key = f"upload_dedup:{file_hash}"
            existing_tid = REDIS.get(key)
            if existing_tid:
                # If we already have a task for this exact file, return it
                t = task_get(existing_tid)
                if t and t.get('status') not in ('error',):
                    return jsonify({"task_id": existing_tid, "dedup": True})
            # else set mapping for short TTL (10 minutes)
            # tid will be set after we create the task below
        except Exception:
            pass

    # Quick validation: ensure it's a readable DOCX
    try:
        from docx import Document as _Doc
        _ = _Doc(in_path)
    except Exception:
        try:
            os.remove(in_path)
        except Exception:
            pass
        return jsonify({"error": "invalid_docx", "message": "Corrupted or unreadable .docx"}), 400

    task_id = str(uuid.uuid4())
    task_create(task_id, {
        "status": "queued",
        "total": 0,
        "done": 0,
        "started_at": time.time(),
        "finished_at": None,
        "filename_in": in_path,
        "filename_out": out_path,
    })

    # Store dedup mapping now that we have a task id
    if USE_REDIS and REDIS is not None:
        try:
            key = f"upload_dedup:{file_hash}"
            # set if not exists, expire after 10 minutes
            REDIS.setex(key, 600, task_id)
        except Exception:
            pass

    # If RQ is enabled, enqueue to Redis-backed queue, else use local thread dispatcher
    if getattr(Config, "USE_RQ_TASKS", False):
        if REDIS is None or RQQueue is None:
            task_merge(task_id, {"status": "error", "finished_at": time.time()})
            return jsonify({"error": "rq_unavailable"}), 500
        try:
            q = RQQueue(getattr(Config, "RQ_QUEUE_NAME", "docjobs"), connection=REDIS)
            # Enqueue worker job by dotted path to avoid import issues
            job = q.enqueue("tasks.rq_process_document_task", task_id, job_timeout=getattr(Config, "RQ_JOB_TIMEOUT", 7200))
            # approximate queue position
            task_merge(task_id, {"queue_pos": q.count})
        except Exception as e:
            task_merge(task_id, {"status": "error", "finished_at": time.time()})
            return jsonify({"error": "enqueue_failed", "message": str(e)}), 500
        return jsonify({"task_id": task_id})

    # Fallback: local bounded queue
    try:
        _PENDING_Q.put_nowait(task_id)
        task_merge(task_id, {"queue_pos": _PENDING_Q.qsize()})
        return jsonify({"task_id": task_id})
    except queue.Full:
        task_merge(task_id, {"status": "rejected", "finished_at": time.time()})
        return jsonify({
            "error": "queue_full",
            "message": "Document queue is full, please retry later.",
            "queue_limit": DOC_QUEUE_MAX
        }), 429

### TODO: testing semaphore: ===>

### TODO: working solutions: <===
    # t = threading.Thread(target=_process_document_task, args=(task_id,), daemon=True)
    # t.start()

    # return jsonify({"task_id": task_id})
### TODO: working solutions: ===>

@app.get("/api/queue_info")
def api_queue_info():
    dprint(f"{'='*33} def api_queue_info(): {'='*33}")
    
    in_progress = None
    # best-effort: _value is CPython-specific; if not present, omit
    if hasattr(_DOC_SEM, "_value"):
        in_progress = DOC_MAX_CONCURRENCY - _DOC_SEM._value
    if getattr(Config, "USE_RQ_TASKS", False) and REDIS is not None and RQQueue is not None:
        try:
            q = RQQueue(getattr(Config, "RQ_QUEUE_NAME", "docjobs"), connection=REDIS)
            return jsonify({
                "mode": "rq",
                "queued": q.count,
                "in_progress": None,
            })
        except Exception:
            pass
    return jsonify({
        "mode": "local",
        "active_limit": DOC_MAX_CONCURRENCY,
        "queued": _PENDING_Q.qsize(),
        "in_progress": in_progress
    })

@app.get("/api/health")
def api_health():
    """Minimal health/status for readiness checks."""
    in_progress = None
    if hasattr(_DOC_SEM, "_value"):
        in_progress = DOC_MAX_CONCURRENCY - _DOC_SEM._value
    return jsonify({
        "status": "ok",
        "ready": all([DB is not None, LLM_CHAT is not None, LLM_DOC is not None]),
        "queue": {
            "active_limit": DOC_MAX_CONCURRENCY,
            "queued": _PENDING_Q.qsize(),
            "in_progress": in_progress,
        },
        "time": datetime.utcnow().isoformat() + "Z",
    })

@app.get("/api/progress/<task_id>")
def api_progress(task_id):
    dprint(f"{'='*33} def api_progress(task_id): {'='*33}")
    
    task = task_get(task_id)
    if not task:
        return abort(404)

    def event_stream():
        # advise client to retry quickly on disconnects
        yield "retry: 5000\n\n"
        # push initial snapshot
        yield f"data: {json.dumps(_task_snapshot(task_id))}\n\n"
        while True:
            if USE_REDIS and REDIS is not None:
                # block on new events with 30s timeout
                res = REDIS.brpop(_task_events_key(task_id), timeout=10)
                if res is not None:
                    _list, payload = res
                    try:
                        msg = json.loads(payload)
                    except Exception:
                        msg = {"status": "unknown"}
                    yield f"data: {json.dumps(msg)}\n\n"
                    if msg.get("status") in ("done", "error"):
                        break
                else:
                    # heartbeat snapshot
                    snap = _task_snapshot(task_id)
                    yield f"data: {json.dumps(snap)}\n\n"
                    if snap.get("status") in ("done", "error"):
                        break
            else:
                # in-memory queue
                q = TASKS.get(task_id, {}).get("stream_queue")
                try:
                    msg = q.get(timeout=10.0)  # heartbeat max
                    yield f"data: {json.dumps(msg)}\n\n"
                    if msg.get("status") in ("done", "error"):
                        break
                except queue.Empty:
                    # periodic ping with snapshot
                    snap = _task_snapshot(task_id)
                    yield f"data: {json.dumps(snap)}\n\n"
                    t = task_get(task_id)
                    if t and t.get("status") in ("done", "error"):
                        break

    # SSE response with no-buffering headers for reverse proxies (e.g., Nginx)
    resp = Response(event_stream(), mimetype="text/event-stream")
    resp.headers["Cache-Control"] = "no-cache"
    resp.headers["X-Accel-Buffering"] = "no"  # Nginx: disable proxy buffering
    resp.headers["Connection"] = "keep-alive"
    return resp


# --- Error handlers ---
@app.errorhandler(413)
def handle_large_file(e):
    return jsonify({"error": "file_too_large", "max_bytes": app.config.get("MAX_CONTENT_LENGTH")}), 413

@app.get("/api/download/<task_id>")
def api_download(task_id):
    dprint(f"{'='*33} def api_download(task_id): {'='*33}")
    
    task = task_get(task_id)
    if not task:
        return abort(404)
    if task["status"] != "done":
        return abort(409)  # not ready
    if not os.path.exists(task["filename_out"]):
        return abort(404)
    return send_file(task["filename_out"], as_attachment=True)

@app.get("/api/results")
def api_results():
    dprint(f"{'='*33} def api_results(): {'='*33}")
    
    """
    List processed (and in-progress) result files from the docx_results folder.
    If a task_id exists for the file, return a download URL via /api/download/<task_id>.
    Otherwise expose a raw download via /api/raw_download?name=<basename>.
    """
    root = Config.RESULT_PATH
    os.makedirs(root, exist_ok=True)
    entries = []
    # Build reverse map only for in-memory; for Redis, use raw download links
    out_to_task = {}
    if not USE_REDIS:
        for tid, t in TASKS.items():
            out_to_task[os.path.abspath(t.get("filename_out",""))] = tid
    for name in sorted(os.listdir(root)):
        if not name.lower().endswith(".docx"):
            continue
        p = os.path.abspath(os.path.join(root, name))
        try:
            st = os.stat(p)
        except FileNotFoundError:
            continue
        task_id = out_to_task.get(p)
        # Prefer /api/download/<task_id> if known and status is done; else raw
        href = None
        if task_id and not USE_REDIS and TASKS.get(task_id, {}).get("status") == "done":
            href = f"/api/download/{task_id}"
        else:
            href = f"/api/raw_download?name={name}"
        entries.append({
            "name": name,
            "size": st.st_size,
            "mtime": int(st.st_mtime),
            "task_id": task_id,
            "href": href,
        })
    # Sort by mtime desc
    entries.sort(key=lambda x: x["mtime"], reverse=True)

    # Optional pagination via query params: page (1-based) and page_size
    page_param = request.args.get("page")
    page_size_param = request.args.get("page_size") or request.args.get("limit")
    if page_param or page_size_param:
        try:
            page = max(int(page_param or 1), 1)
        except Exception:
            page = 1
        try:
            default_ps = getattr(Config, "RESULTS_PAGE_SIZE_DEFAULT", 50)
            max_ps = getattr(Config, "RESULTS_PAGE_SIZE_MAX", 200)
            page_size = max(min(int(page_size_param or default_ps), max_ps), 1)
        except Exception:
            page_size = getattr(Config, "RESULTS_PAGE_SIZE_DEFAULT", 50)
        total = len(entries)
        total_pages = (total + page_size - 1) // page_size
        start = (page - 1) * page_size
        end = start + page_size
        sliced = entries[start:end]
        return jsonify({
            "results": sliced,
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": total_pages,
        })

    # Back-compat: if no paging params, return full list
    return jsonify({"results": entries})

@app.get("/api/raw_download")
def api_raw_download():
    dprint(f"{'='*33} def api_raw_download(): {'='*33}")
    
    """
    Directly download a file by name from docx_results.
    Used as a fallback when we do not have a task_id (e.g., after server restart).
    """
    name = request.args.get("name", "").strip()
    if not name or "/" in name or "\\" in name:
        return abort(400)
    root = "docx_results"
    path = os.path.join(root, name)
    if not os.path.exists(path):
        return abort(404)
    return send_file(path, as_attachment=True)

# ---------- Corpus (documents/) management ----------
@app.get("/api/corpus/list")
@limiter.limit("30/minute")
def api_corpus_list():
    root = Config.DOCS_PATH
    os.makedirs(root, exist_ok=True)

    # Optional pagination
    try:
        page = max(int(request.args.get("page", "1")), 1)
    except Exception:
        page = 1
    try:
        page_size = max(min(int(request.args.get("page_size", "200")), 1000), 1)
    except Exception:
        page_size = 200

    entries = []
    md_root_abs = os.path.abspath(getattr(Config, 'DOCS_PATH_MD', os.path.join(root, 'documents_converted_to_md')))
    root_abs = os.path.abspath(root)
    for dirpath, dirnames, files in os.walk(root):
        # prune converted MD subtree so it doesn't appear in UI
        try:
            dirnames[:] = [d for d in dirnames if not os.path.abspath(os.path.join(dirpath, d)).startswith(md_root_abs)]
        except Exception:
            pass
        # skip files under the MD subtree
        if os.path.abspath(dirpath).startswith(md_root_abs):
            continue
        # Add directories (so empty folders show up)
        for d in dirnames:
            full_dir = os.path.join(dirpath, d)
            try:
                st = os.stat(full_dir)
            except FileNotFoundError:
                continue
            rel_dir = os.path.relpath(full_dir, root)
            entries.append({
                "relpath": rel_dir,
                "mtime": int(st.st_mtime),
                "kind": "dir",
            })
        # Add files
        for name in files:
            if not name.lower().endswith('.docx'):
                continue
            full = os.path.join(dirpath, name)
            try:
                st = os.stat(full)
            except FileNotFoundError:
                continue
            rel = os.path.relpath(full, root)
            entries.append({
                "relpath": rel,
                "size": st.st_size,
                "mtime": int(st.st_mtime),
                "kind": "file",
            })
    entries.sort(key=lambda x: (x["relpath"]))

    total = len(entries)
    start = (page - 1) * page_size
    end = start + page_size
    slice_entries = entries[start:end]
    return jsonify({
        "root": root,
        "total": total,
        "page": page,
        "page_size": page_size,
        "results": slice_entries,
    })


@app.post("/api/corpus/upload")
@limiter.limit("10/minute")
def api_corpus_upload():
    if "file" not in request.files:
        return jsonify({"error": "no file"}), 400
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"error": "no filename"}), 400
    orig_name = f.filename or ""
    base_name = os.path.basename(orig_name)
    if not _is_docx_filename(base_name):
        return jsonify({"error": "only .docx is supported"}), 400

    rel_dir = request.form.get("rel_dir", "").strip().strip("/\\")
    root = Config.DOCS_PATH
    os.makedirs(root, exist_ok=True)
    dest_dir = os.path.join(root, rel_dir) if rel_dir else root
    os.makedirs(dest_dir, exist_ok=True)
    dest = os.path.abspath(os.path.join(dest_dir, base_name))
    # path traversal guard: ensure under root
    if not dest.startswith(os.path.abspath(root) + os.sep) and dest != os.path.abspath(root):
        return jsonify({"error": "invalid_path"}), 400

    f.save(dest)
    try:
        st = os.stat(dest)
    except Exception:
        st = None
    return jsonify({
        "relpath": os.path.relpath(dest, root),
        "size": getattr(st, 'st_size', None),
        "mtime": int(getattr(st, 'st_mtime', time.time())),
        "saved": True,
    })


@app.post("/api/corpus/delete")
@limiter.limit("10/minute")
def api_corpus_delete():
    data = request.get_json(silent=True) or {}
    relpath = (data.get("relpath") or "").strip().strip("/\\")
    if not relpath:
        return jsonify({"error": "missing_relpath"}), 400
    root = Config.DOCS_PATH
    full = os.path.abspath(os.path.join(root, relpath))
    # ensure under root
    if not full.startswith(os.path.abspath(root) + os.sep):
        return jsonify({"error": "invalid_path"}), 400
    if not os.path.exists(full):
        return jsonify({"error": "not_found"}), 404
    if not full.lower().endswith('.docx'):
        return jsonify({"error": "invalid_extension"}), 400

    # Optional: also remove vectors for this IRD (prefer content-hash before unlink)
    try:
        from rag_logic import _open_db_for_update
        import hashlib
        # Prefer content hash for deletion (compute before file unlink)
        def _sha256_file(p: str):
            try:
                h = hashlib.sha256()
                with open(p, 'rb') as f:
                    for chunk in iter(lambda: f.read(1024*1024), b''):
                        h.update(chunk)
                return h.hexdigest()
            except Exception:
                return None
        group_id = _sha256_file(full)
        db = _open_db_for_update()
        # delete content-hash id first
        if group_id:
            try:
                db.delete(filter={"attachment_group_id": group_id})
            except Exception:
                pass
        # delete legacy ids
        try:
            legacy_gid = hashlib.sha1(os.path.splitext(os.path.basename(full))[0].encode('utf-8')).hexdigest()
            db.delete(filter={"attachment_group_id": legacy_gid})
        except Exception:
            pass
        try:
            rel = os.path.relpath(full, Config.DOCS_PATH)
            rel_no_ext = os.path.splitext(rel.replace("\\", "/"))[0]
            legacy_path_gid = hashlib.sha1(rel_no_ext.encode('utf-8')).hexdigest()
            db.delete(filter={"attachment_group_id": legacy_path_gid})
        except Exception:
            pass
        # 4) Final safety net: delete by docx_relpath metadata match
        try:
            rel_norm = os.path.relpath(full, Config.DOCS_PATH).replace("\\", "/")
            db.delete(filter={"docx_relpath": rel_norm})
        except Exception:
            pass
        # PersistentClient writes to disk automatically; no explicit persist needed
    except Exception:
        pass

    # Best-effort: delete file from filesystem (after vectors removed)
    try:
        os.remove(full)
    except Exception as e:
        return jsonify({"error": "delete_failed", "message": str(e)}), 500

    return jsonify({"deleted": True, "relpath": relpath})


@app.post("/api/corpus/mkdir")
@limiter.limit("20/minute")
def api_corpus_mkdir():
    """Create a folder under documents/ (optionally inside a parent)."""
    data = request.get_json(silent=True) or {}
    parent = (data.get("parent") or "").strip().strip("/\\")
    name = (data.get("name") or "").strip().strip("/\\")
    if not name:
        return jsonify({"error": "missing_name"}), 400
    # basic name guard
    if any(ch in name for ch in ("/", "\\")):
        return jsonify({"error": "invalid_name"}), 400
    root = Config.DOCS_PATH
    base = os.path.join(root, parent) if parent else root
    dest = os.path.abspath(os.path.join(base, name))
    # path traversal guard: ensure under root
    if not dest.startswith(os.path.abspath(root) + os.sep):
        return jsonify({"error": "invalid_path"}), 400
    try:
        os.makedirs(dest, exist_ok=True)
        st = os.stat(dest)
        rel = os.path.relpath(dest, root)
        return jsonify({
            "created": True,
            "relpath": rel,
            "mtime": int(getattr(st, 'st_mtime', time.time())),
        })
    except Exception as e:
        return jsonify({"error": "mkdir_failed", "message": str(e)}), 500


@app.post("/api/corpus/rmdir")
@limiter.limit("10/minute")
def api_corpus_rmdir():
    """Recursively remove a folder under documents/ and its vectors for contained .docx files."""
    import shutil
    data = request.get_json(silent=True) or {}
    relpath = (data.get("relpath") or "").strip().strip("/\\")
    if relpath in ("", "."):
        return jsonify({"error": "refuse_remove_root"}), 400
    root = Config.DOCS_PATH
    full = os.path.abspath(os.path.join(root, relpath))
    if not full.startswith(os.path.abspath(root) + os.sep):
        return jsonify({"error": "invalid_path"}), 400
    if not os.path.exists(full) or not os.path.isdir(full):
        return jsonify({"error": "not_found_or_not_dir"}), 404

    # collect all .docx in subtree to delete vectors
    docx_list = []
    for dirpath, _dirs, files in os.walk(full):
        for name in files:
            if name.lower().endswith('.docx'):
                docx_list.append(os.path.join(dirpath, name))
    # delete vectors best-effort
    try:
        from rag_logic import _open_db_for_update
        import hashlib
        if docx_list:
            db = _open_db_for_update()
            for p in docx_list:
                def _sha256_file(pth: str):
                    try:
                        h = hashlib.sha256()
                        with open(pth, 'rb') as f:
                            for chunk in iter(lambda: f.read(1024*1024), b''):
                                h.update(chunk)
                        return h.hexdigest()
                    except Exception:
                        return None
                group_id = _sha256_file(p)
                try:
                    if group_id:
                        db.delete(where={"attachment_group_id": group_id})
                except Exception:
                    pass
                try:
                    legacy_gid = hashlib.sha1(os.path.splitext(os.path.basename(p))[0].encode('utf-8')).hexdigest()
                    db.delete(where={"attachment_group_id": legacy_gid})
                except Exception:
                    pass
                try:
                    rel = os.path.relpath(p, Config.DOCS_PATH)
                    rel_no_ext = os.path.splitext(rel.replace("\\", "/"))[0]
                    legacy_path_gid = hashlib.sha1(rel_no_ext.encode('utf-8')).hexdigest()
                    db.delete(where={"attachment_group_id": legacy_path_gid})
                except Exception:
                    pass
            # try:
            #     # No explicit persist required with PersistentClient backend

            # except Exception:
            #     pass
    except Exception:
        pass

    # remove the folder recursively
    try:
        shutil.rmtree(full)
    except Exception as e:
        return jsonify({"error": "rmdir_failed", "message": str(e)}), 500

    return jsonify({"deleted": True, "relpath": relpath})

def _task_snapshot(task_id):
    dprint(f"{'='*33} def _task_snapshot(task_id): {'='*33}")
    
    t = task_get(task_id)
    if not t:
        return {"task_id": task_id, "status": "unknown"}
    return {
        "task_id": task_id,
        "status": t["status"],
        "total": t["total"],
        "done": t["done"],
        "filename_in": os.path.basename(t["filename_in"]),
        "ready": (t["status"] == "done"),
    }


# ---------- Worker ----------
def _process_document_task(task_id: str):
    dprint(f"{'='*33} def _process_document_task(task_id: str): {'='*33}")
    
    task = task_get(task_id)
    if not task:
        return

    try:
        task_merge(task_id, {"status": "running"})
        task_push_event(task_id, _task_snapshot(task_id))

        nbu_file = task.get("filename_in")

        # Split NBU doc into questions (with para ids)
        buckets = split_docx_to_question_with_ids(nbu_file, second_split=True)
        total = len(buckets)
        task_merge(task_id, {"total": total})
        task_push_event(task_id, _task_snapshot(task_id))

        # Prepare dense retriever (we'll derive sparse BM25 over dense candidates per‑subquery)
        dense = DB.as_retriever(
            search_kwargs={
                "k": getattr(Config, "HYBRID_DENSE_K", 12),
                "fetch_k": getattr(Config, "HYBRID_DENSE_K", 12) * 3,
            },
            search_type="mmr",
        )

        from docx import Document
        doc = Document(nbu_file)  # open copy to annotate
        done = 0

        # Prompt skeleton (JSON expected by your system prompt)
        system_prompt = Config.system_prompt_document_loop

        for b in buckets:
            q_text = b["question_text"]
            para_indices0 = b.get("para_indices0", []) or b.get("para_ids", [])
            if not isinstance(para_indices0, list):
                para_indices0 = []

            # Build fused, capped, packed context via shared helper
            packed, _ = build_packed_context(DB, q_text)

            context_str = "\n\n".join(
                f"===\n[Джерело: {_source_label_from_meta(d.metadata)}]\n{d.page_content}\n===" for d in packed
            )

            prompt = f"""{system_prompt}
Контекст з ВНД:
{context_str}

Вимога НБУ:
{q_text}

Відповідь українською:
"""

            # LLM call (AzureChatOpenAI with response_format="json_object")
            resp = _invoke_with_backoff(LLM_DOC, prompt)
            dprint(f"resp: {resp}")
            raw = resp.content if hasattr(resp, "content") else str(resp)
            dprint(f"raw: {raw}")

            # parse JSON (your example_json contract)
            verdict = "❓"; text_answer = raw; sources = []
            try:
                data = json.loads(raw)
                text_answer = data.get("text", raw)
                dprint(f"text_answer : {text_answer}")
                sources = data.get("source", [])
                verdict = data.get("answer", "❓")
                if verdict not in ("✅", "❌", "❓"):
                    verdict = "❓"
                # Strip stray NULs and normalize line breaks coming from the model
                verdict = (verdict or "").replace("\x00", "")
                text_answer = sanitize_markdown(text_answer)
            except Exception:
                pass  # if LLM returns non-JSON, still continue

            color = _color_from_answer(verdict)

            # Write a comment into the referenced paragraph range
            # b['para_indices0'] are zero-based indices as produced by your splitter
            if para_indices0:
                try:
                    comment_text = (
                        f"{verdict} {text_answer.strip()}\n\nДжерела:\n" +
                        ("\n".join(f"- {s}" for s in sources) if isinstance(sources, list) else str(sources))
                    )

                    dprint(f"comment_text : ")
                    dprint(comment_text)

                    add_comment_to_paragraphs(
                        doc,
                        para_indices0,
                        comment_text,
                        author="RAG Assistant",
                        initials="AI",
                        color=color
)
                    # add_comment_to_paragraphs(
                    #     doc,
                    #     para_indices0,
                    #     f"{verdict} {text_answer}\nДжерела: {', '.join(sources) if isinstance(sources, list) else str(sources)}",
                    #     author="RAG Assistant",
                    #     initials="RA",
                    #     color=color
                    # )
                except Exception as e:
                    # keep going even if comment fails
                    pass

            done += 1
            task_merge(task_id, {"done": done})

            # push incremental progress
            task_push_event(task_id, {
                **_task_snapshot(task_id),
                "last_verdict": verdict,
                "last_brief": (text_answer[:140] + "…") if len(text_answer) > 140 else text_answer
            })

        # Save annotated DOCX
        out_path = task["filename_out"]
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)

        # Record processed file mapping for historical dedup
        try:
            ihash = sha256_file(task.get("filename_in", ""))
            if ihash:
                cache_store(ihash, out_path, meta={"task_id": task_id})
        except Exception:
            pass

        task_merge(task_id, {"status": "done", "finished_at": time.time()})
        task_push_event(task_id, _task_snapshot(task_id))

    except Exception as e:
        task_merge(task_id, {"status": "error", "finished_at": time.time()})
        task_push_event(task_id, {"status": "error", "error": str(e), **_task_snapshot(task_id)})


@app.post("/api/corpus/sync")
@limiter.limit("5/minute")
def api_corpus_sync():
    """Trigger vector store sync from documents/.

    Body (JSON, optional): {"mode": "apply"|"dry"|"init"}
    Default mode is "apply".
    """
    body = request.get_json(silent=True) or {}
    mode = (body.get("mode") or "apply").lower()
    init = (mode == "init")
    dry = (mode == "dry")
    try:
        from scripts.sync_vectorstore import run_sync, _default_manifest_path, run_init_from_db
        manifest = _default_manifest_path()
        # If manifest is missing and caller requested apply, do a safe init instead
        if not os.path.exists(manifest) and not init and not dry and mode != 'init_db':
            init = True
        if mode == 'init_db':
            report = run_init_from_db(manifest_path=manifest)
            status = 'initialized_from_db'
        else:
            report = run_sync(manifest_path=manifest, init_manifest=init, dry_run=dry)
            status = "initialized" if init else ("dry" if dry else "applied")
        return jsonify({"ok": True, "status": status, "report": report})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/api/corpus/verify")
@limiter.limit("10/minute")
def api_corpus_verify():
    try:
        from scripts.sync_vectorstore import run_verify
        report = run_verify()
        return jsonify({"ok": True, "report": report})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/source_label', methods=['GET', 'POST'])
def api_source_label():
    """Get or set the source label mode used in prompts.

    GET:  {"mode": "filename"|"docx_relpath"}
    POST: {"mode": "filename"|"docx_relpath"}
    """
    global SOURCE_LABEL_MODE
    if request.method == 'GET':
        return jsonify({"mode": SOURCE_LABEL_MODE})
    data = request.get_json(silent=True) or {}
    mode = (data.get('mode') or '').strip().lower()
    if mode not in ('filename', 'docx_relpath'):
        return jsonify({"error": "invalid_mode"}), 400
    SOURCE_LABEL_MODE = mode
    return jsonify({"ok": True, "mode": SOURCE_LABEL_MODE})

# Source label mode (runtime-adjustable for the web process)
SOURCE_LABEL_MODE = getattr(Config, 'SOURCE_LABEL_MODE', 'filename').strip().lower()

def _source_label_from_meta(meta: dict) -> str:
    return source_label_from_meta(meta, mode=SOURCE_LABEL_MODE)


# ---------- Main ----------
if __name__ == "__main__":
    dprint(f"{'='*33} if __name__ == __main__: app.py {'='*33}")
    # For dev
    app.run(host="0.0.0.0", port=20000, debug=DEBUG, threaded=True)
