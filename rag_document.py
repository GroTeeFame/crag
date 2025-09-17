import os
import re
import time
import json
import datetime
import threading

from langchain.schema import Document as LangchainDocument
from langchain.prompts import PromptTemplate
from docx import Document

from colorama import Fore, Style  # Terminal colors

from config import Config

from read_split_nbu import split_docx_to_question_with_ids

from save_results import add_comment_to_paragraphs

from rag_functions import (
    count_tokens, spinner, wait_if_tpm_exceeded, question_giver,
    build_packed_context, source_label_from_meta
)

from collections import deque
token_usage_log = deque()

window_start_time = time.time()

# Debug gating via Config.DEBUG
DEBUG = getattr(Config, "DEBUG", False)

def dprint(msg: str):
    if DEBUG:
        print(msg)


def iter_docs_from_chroma(db, page_size: int = 1000):
    """Yield LangchainDocument objects from Chroma in pages to keep memory bounded."""
    col = db._collection
    try:
        total = col.count()
        offset = 0
        while offset < total:
            res = col.get(include=["documents", "metadatas"], limit=page_size, offset=offset)
            docs = res.get("documents", []) or []
            metas = res.get("metadatas", []) or []
            got = 0
            for content, meta in zip(docs, metas):
                got += 1
                yield LangchainDocument(page_content=content, metadata=meta or {})
            if not got:
                break
            ids = res.get("ids", [])
            offset += len(ids) if ids else got
    except Exception:
        # Fallback by IDs
        id_res = col.get(include=[])
        ids = id_res.get("ids", []) if isinstance(id_res, dict) else []
        for i in range(0, len(ids), page_size):
            sl = ids[i:i+page_size]
            res = col.get(ids=sl, include=["documents", "metadatas"])
            for content, meta in zip(res.get("documents", []) or [], res.get("metadatas", []) or []):
                yield LangchainDocument(page_content=content, metadata=meta or {})


def _make_subqueries(q: str) -> list[str]:
    parts = [x.strip() for x in q.splitlines() if x.strip()]
    out = []
    for line in parts:
        if re.match(r'^\s*(?:\d+(\.\d+)*\.|\-|\•)\s+', line):
            out.append(re.sub(r'^\s*(?:\d+(\.\d+)*\.|\-|\•)\s+', '', line))
        else:
            out.append(line)
    if q not in out:
        out.append(q)
    seen = set(); dedup = []
    for s in out:
        if s not in seen:
            dedup.append(s); seen.add(s)
    return dedup


def document_logic(db, llm, nbu_document_name):  # TODO: handle empty NBU gracefully
    dprint(f"{'='*33} def document_logic(db, llm, nbu_document_name) {'='*33}")


    dprint(Fore.LIGHTRED_EX + f"inside document_logic" + '-'*70 + Fore.RESET)

    system_prompt = Config.system_prompt_document_loop
    
    prompt_template = PromptTemplate(
        template=f"""
        {system_prompt}
        Контекст з ВНД:
        {{context}}

        Вимога НБУ:
        {{question}}

        Відповідь українською:
    """,
        input_variables=["context", "question"]
    )

    nbu_file_path = os.path.join(Config.NBU_DOCS_PATH, nbu_document_name)

    dprint(Fore.RED + nbu_file_path + Fore.RESET)

    # Prepare questions extracted from NBU doc
    question_from_nbu = split_docx_to_question_with_ids(nbu_file_path, True)  # with ids, json
    if not question_from_nbu:
        print(Fore.RED + "No questions extracted from NBU document. Exiting." + Fore.RESET)
        return

    ct = datetime.datetime.now()
    fct = ct.strftime("%d-%m-%Y %H-%M-%S")

    os.makedirs('full_prompt', exist_ok=True)
    full_prompt_txt_filename = f"full_prompt/full_prompt_from_result_of_NBU({nbu_document_name[:-5]})_to_VND_comparison_{fct}_DN({Config.DEPLOYMENT_NAME}).txt"

    if Config.CREATE_NEW_RESULT_TXT:
        os.makedirs('work_results', exist_ok=True)
        result_txt_filename = f"work_results/result_of_NBU({nbu_document_name[:-5]})_to_VND_comparison_{fct}_DN({Config.DEPLOYMENT_NAME}).txt"
        questions_generator = question_giver(question_from_nbu)
    else:
        result_txt_filename = Config.CONTINUE_RESULT_TXT_NAME
        questions_generator = question_giver(question_from_nbu)
        
    doc = Document(nbu_file_path)


    # Retrieval is handled by shared helper now


    with open(result_txt_filename, "a") as txt_file_to_write_result:
        while True:

            ct = datetime.datetime.now()
            fct = ct.strftime("%d-%m-%Y %H:%M:%S")

            try:
                question_obj = next(questions_generator)
                question = question_obj.copy()

            except StopIteration:
                print(Fore.GREEN + "✅ All questions processed. Exiting loop." + Fore.RESET)
                os.makedirs('docx_results', exist_ok=True)
                doc.save(f"docx_results/result_of_NBU({nbu_document_name[:-5]})_to_VND_comparison_{fct}_DN({Config.DEPLOYMENT_NAME}).docx")
                break

            # Pull human-readable text of the question
            query = question["question_text"]

            splited_query = query.splitlines()
            
            dprint(Fore.YELLOW + f" Query was splitted to {len(splited_query)} pieces" + Fore.RESET)

            retrieved_chunks = Config.RETURN_CHUNK

#=========

            full_context = []

            # Build packed context for entire query via shared helper
            packed, _ = build_packed_context(db, query)
            full_context_str = "\n\n".join(
                f"===\n[Джерело: {source_label_from_meta(doc.metadata)}]\n{doc.page_content}\n==="
                for doc in packed
            )

#=========

            # Estimate token count from context + question
            token_count = count_tokens(full_context_str) + count_tokens(query)

            formatted_prompt = prompt_template.format(context=full_context_str, question=query)


            stop_spinner = threading.Event()
            spinner_thread = threading.Thread(target=spinner, args=("Thinking...", stop_spinner))

            estimated_total_tokens = token_count + 1000
            wait_if_tpm_exceeded(estimated_total_tokens)
            spinner_thread.start()
            result = None
            try:  ### Call LLM
                # print("here must be calling to GPT!!!")
                # result = ["ITS ONLY NOTHING!!!"]
                token_usage_log.append((time.time(), estimated_total_tokens))  # Temporarily log estimate
                dprint(Fore.CYAN + f"🔢 Estimated prompt tokens: {token_count}" + Fore.RESET)
                
                result = llm.invoke(formatted_prompt)  # current working LLM call
                token_usage_log.pop()  # Remove estimated token usage
                used_tokens = result.response_metadata.get('token_usage', {}).get('total_tokens', estimated_total_tokens)
                token_usage_log.append((time.time(), used_tokens))

            finally:
                stop_spinner.set()
                spinner_thread.join()

            dprint("\n--- Question: ---")
            dprint(Fore.LIGHTMAGENTA_EX + Style.DIM + query + Fore.RESET + Style.RESET_ALL)

            dprint("\n--- Answer: ---")
            dprint(Fore.GREEN)

            if result is None:
                print(Fore.RED + "LLM invocation failed; skipping this question." + Fore.RESET)
                txt_file_to_write_result.write(f"Question: \n{question['question_text']}\n\n")
                txt_file_to_write_result.write("Answer from GPT: \n[ERROR: LLM invocation failed]\n")
                txt_file_to_write_result.write(f"-------------------------------------------------\n")
                continue
            
            if hasattr(result, "content"):
                try:
                    answer = json.loads(result.content)
                except Exception as e:
                    print(Fore.RED + f"Failed to parse JSON response: {e}" + Fore.RESET)
                    print(Fore.YELLOW + f"Raw content: {result.content}" + Fore.RESET)
                    answer = {"text": "", "source": [], "answer": "❓"}
# json.loads(llm.invoke(prompt).content)

                # answer = parse_gpt_json_output(result.content)
                
                dprint(f"answer['text']: {answer['text']}")
                dprint(f"answer['source']: {answer['source']}")
                dprint(f"answer['answer']: {answer['answer']}")

                with open(full_prompt_txt_filename, "a") as full_prompt_txt_file_to_write_result:
                    full_prompt_txt_file_to_write_result.write(formatted_prompt)
                    full_prompt_txt_file_to_write_result.write("\n--=============================================================\n")
                    full_prompt_txt_file_to_write_result.write(f"answer['text']: {answer['text']}\n")
                    full_prompt_txt_file_to_write_result.write(f"answer['source']: {answer['source']}\n")
                    full_prompt_txt_file_to_write_result.write(f"answer['answer']: {answer['answer']}\n")
                    full_prompt_txt_file_to_write_result.write("===============================================================\n\n")


                comment = f"{answer['text']} | Джерела: {answer['source']}"

                add_comment_to_paragraphs(doc, question_obj["para_indices0"], comment, author="AI helper", color=answer["answer"]) # new func with para comment

                token_usage = result.response_metadata.get('token_usage', {})
                # for item in result.response_metadata['token_usage']:
                dprint(Fore.YELLOW + f"---------------------------")
                if DEBUG and token_usage:
                    print(f"completion_tokens : {token_usage.get('completion_tokens')}")
                    print(f"prompt_tokens : {token_usage.get('prompt_tokens')}")
                    print(f"total_tokens : {token_usage.get('total_tokens')}")
                elif DEBUG:
                    print("Token usage metadata is not available")
                dprint(f"---------------------------" + Fore.RESET)
                dprint(Fore.CYAN + f"---------------------------")
                dprint(f"question was splitted to {len(splited_query)} question")
                # dprint(f"for each question was used {calculated_chunk_buffer} chunks")
                dprint(f"for each question was used {retrieved_chunks} chunks")
                dprint(f"---------------------------" + Fore.RESET)

            else:
                dprint(str(result))
            dprint(Fore.RESET)


            txt_file_to_write_result.write(f"Question: \n{question['question_text']}\n\n")
            txt_file_to_write_result.write(f"Answer from GPT: \n{result.content}\n")
            txt_file_to_write_result.write(f"-------------------------------------------------\n")

            elapsed_time = time.time() - window_start_time
            formatted_time = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
            dprint(Fore.MAGENTA + f"⏱️ Script running time: {formatted_time}" + Fore.RESET)
