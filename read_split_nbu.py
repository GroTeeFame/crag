import os
import sys
import re
import json
from typing import List, Dict, Optional
from collections import defaultdict

from colorama import Fore

from docx import Document
from docx.oxml.ns import qn

# Debug gating via Config.DEBUG
try:
    from config import Config
    DEBUG = getattr(Config, "DEBUG", False)
except Exception:
    DEBUG = False

def dprint(msg: str):
    if DEBUG:
        print(msg)


def _to_letter(n: int, uppercase: bool = True) -> str:
    if n <= 0:
        return ""
    result = []
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        base = ord('A') if uppercase else ord('a')
        result.append(chr(base + remainder))
    return ''.join(reversed(result))


def _to_roman(n: int, uppercase: bool = True) -> str:
    if n <= 0:
        return ""
    numerals = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    result = []
    for value, symbol in numerals:
        while n >= value:
            result.append(symbol)
            n -= value
    roman = ''.join(result)
    return roman if uppercase else roman.lower()


def _format_number_value(value: int, fmt: str) -> str:
    fmt_lower = (fmt or 'decimal').lower()
    if fmt_lower in {'decimal', 'decimalzero', 'ordinal', 'ordinaltext', 'cardinaltext'}:
        return str(value)
    if fmt_lower == 'upperletter':
        return _to_letter(value, uppercase=True)
    if fmt_lower == 'lowerletter':
        return _to_letter(value, uppercase=False)
    if fmt_lower == 'upperroman':
        return _to_roman(value, uppercase=True)
    if fmt_lower == 'lowerroman':
        return _to_roman(value, uppercase=False)
    # Fallback: use decimal representation
    return str(value)


def _build_numbering_cache(doc: Document) -> Optional[Dict[int, Dict[int, Dict[str, object]]]]:
    numbering_part = getattr(doc.part, "numbering_part", None)
    if numbering_part is None:
        return None

    root = numbering_part.element

    num_to_abs: Dict[int, int] = {}
    overrides: Dict[int, Dict[int, int]] = defaultdict(dict)

    for num in root.findall(qn('w:num')):
        num_id_attr = num.get(qn('w:numId'))
        if num_id_attr is None:
            continue
        try:
            num_id = int(num_id_attr)
        except ValueError:
            continue
        abstract = num.find(qn('w:abstractNumId'))
        if abstract is not None and abstract.get(qn('w:val')):
            try:
                num_to_abs[num_id] = int(abstract.get(qn('w:val')))
            except ValueError:
                pass
        for override in num.findall(qn('w:lvlOverride')):
            ilvl_attr = override.get(qn('w:ilvl'))
            if ilvl_attr is None:
                continue
            try:
                ilvl = int(ilvl_attr)
            except ValueError:
                continue
            start_override = override.find(qn('w:startOverride'))
            if start_override is not None and start_override.get(qn('w:val')):
                try:
                    overrides[num_id][ilvl] = int(start_override.get(qn('w:val')))
                except ValueError:
                    pass
            lvl_override = override.find(qn('w:lvl'))
            if lvl_override is not None:
                start_elem = lvl_override.find(qn('w:start'))
                if start_elem is not None and start_elem.get(qn('w:val')):
                    try:
                        overrides[num_id][ilvl] = int(start_elem.get(qn('w:val')))
                    except ValueError:
                        pass

    abstract_defs: Dict[int, Dict[int, Dict[str, object]]] = {}
    for absnum in root.findall(qn('w:abstractNum')):
        abs_id_attr = absnum.get(qn('w:abstractNumId'))
        if abs_id_attr is None:
            continue
        try:
            abs_id = int(abs_id_attr)
        except ValueError:
            continue
        lvl_map: Dict[int, Dict[str, object]] = {}
        for lvl in absnum.findall(qn('w:lvl')):
            ilvl_attr = lvl.get(qn('w:ilvl'))
            if ilvl_attr is None:
                continue
            try:
                ilvl = int(ilvl_attr)
            except ValueError:
                continue
            fmt_elem = lvl.find(qn('w:numFmt'))
            fmt_val = fmt_elem.get(qn('w:val')) if fmt_elem is not None else 'decimal'
            text_elem = lvl.find(qn('w:lvlText'))
            text_val = text_elem.get(qn('w:val')) if text_elem is not None else f"%{ilvl + 1}."
            start_elem = lvl.find(qn('w:start'))
            start_val = 1
            if start_elem is not None and start_elem.get(qn('w:val')):
                try:
                    start_val = int(start_elem.get(qn('w:val')))
                except ValueError:
                    start_val = 1
            lvl_map[ilvl] = {
                "format": fmt_val,
                "text": text_val,
                "start": start_val,
            }
        abstract_defs[abs_id] = lvl_map

    levels: Dict[int, Dict[int, Dict[str, object]]] = {}
    for num_id, abs_id in num_to_abs.items():
        base = abstract_defs.get(abs_id, {})
        merged: Dict[int, Dict[str, object]] = {}
        for ilvl, info in base.items():
            merged[ilvl] = dict(info)
            if ilvl in overrides.get(num_id, {}):
                merged[ilvl]["start"] = overrides[num_id][ilvl]
        for ilvl, start_override in overrides.get(num_id, {}).items():
            if ilvl not in merged:
                merged[ilvl] = {
                    "format": "decimal",
                    "text": f"%{ilvl + 1}.",
                    "start": start_override,
                }
        if merged:
            levels[num_id] = merged

    return levels or None


def _get_numbering_prefix(paragraph, numbering_cache, state) -> str:
    if not numbering_cache:
        return ""
    num_pr = paragraph._element.xpath('./w:pPr/w:numPr')
    if not num_pr:
        return ""
    num_pr = num_pr[0]
    num_id_elem = num_pr.find(qn('w:numId'))
    ilvl_elem = num_pr.find(qn('w:ilvl'))
    if num_id_elem is None or ilvl_elem is None:
        return ""
    try:
        num_id = int(num_id_elem.get(qn('w:val')))
    except (TypeError, ValueError):
        return ""
    try:
        ilvl = int(ilvl_elem.get(qn('w:val')))
    except (TypeError, ValueError):
        ilvl = 0

    level_defs = numbering_cache.get(num_id)
    if not level_defs:
        return ""
    lvl_info = level_defs.get(ilvl)
    if not lvl_info:
        return ""

    num_fmt = (lvl_info.get("format") or "").lower()
    if num_fmt in {"bullet", "none"}:
        return ""

    pattern = lvl_info.get("text", f"%{ilvl + 1}.")
    counters = state.setdefault(num_id, [0] * 10)
    if ilvl >= len(counters):
        counters.extend([0] * (ilvl + 1 - len(counters)))

    for lower in range(ilvl):
        if counters[lower] == 0:
            parent_info = level_defs.get(lower, {})
            counters[lower] = parent_info.get("start", 1)

    if counters[ilvl] == 0:
        counters[ilvl] = lvl_info.get("start", 1)
    else:
        counters[ilvl] += 1

    for deeper in range(ilvl + 1, len(counters)):
        counters[deeper] = 0

    label = pattern
    for idx in range(ilvl + 1):
        lvl_def = level_defs.get(idx, {})
        fmt = lvl_def.get("format", "decimal")
        val = counters[idx] if counters[idx] else lvl_def.get("start", 1)
        formatted = _format_number_value(val, fmt)
        label = label.replace(f"%{idx + 1}", formatted)

    label = re.sub(r'%\d', '', label).strip()
    return label


def normalize_multiline_enumeration(text: str) -> str:
    # dprint("normalize_multiline_enumeration()")

    # Крок 1: перетворити весь текст на список рядків
    lines = text.strip().splitlines()

    # Крок 2: проходимо по рядках, формуючи "поточний пункт"
    result_lines = []
    current = ''

    for line in lines:
        line = line.strip()

        if re.match(r'^\d+\)', line):  # новий пункт 1), 2), ..., 10)
            if current:
                result_lines.append(current.strip())
            current = line
        elif re.match(r'^\d+\.', line):  # наприклад: 117.
            if current:
                result_lines.append(current.strip())
            current = line
        else:
            current += ' ' + line  # додаємо до поточного пункту

    if current:
        result_lines.append(current.strip())

    # Крок 3: обʼєднати всі в один текст
    return '\n'.join(result_lines)

#TODO: доки в тексті не буде знайдено "I. Загальні положення" ми не записуємо текст
#TODO: потрібно знайти фінальну точку документу, після якої буде зупинено запис доку

# --- Paragraph ID helpers ---
def get_paragraph_id(para, fallback_idx: Optional[int] = None) -> str:
    # dprint("get_paragraph_id()")

    """Return Word's paragraph ID (w14:paraId) if present, else a stable fallback.
    The fallback uses the paragraph index if provided, otherwise a hash of text.
    """
    el = para._element
    pid = el.get(qn('w14:paraId')) or el.get(qn('w:paraId'))
    if pid:
        return pid
    # Fallback: deterministic id from index if available, else from text
    if fallback_idx is not None:
        return f"p{fallback_idx:06d}"
    # last resort – avoid very long ids
    return f"p_{abs(hash(para.text)) % (10**12)}"

# --- Extract text with paragraph IDs ---
def extract_text_with_font_and_ids(docx_path: str, not_nbu: bool) -> List[Dict]:
    dprint("extract_text_with_font_and_ids()")

    doc = Document(docx_path)
    numbering_cache = _build_numbering_cache(doc)
    numbering_state: Dict[int, List[int]] = defaultdict(lambda: [0] * 10) if numbering_cache else None
    items = []  # list of {text, para_id}
    start_write_flag = False
    if not_nbu:
        start_write_flag = True

    for i, para in enumerate(doc.paragraphs, 1):
        paragraph_text = ' '.join(para.text.split())
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            if text.find('I. Загальні положення') != -1:
                start_write_flag = True
                paragraph_text = ''
            if text.startswith('{') and text.endswith('}'):
                # skip placeholders
                continue

        if start_write_flag and paragraph_text:
            if paragraph_text.startswith('{') and paragraph_text.endswith('}'):
                continue
            if numbering_cache:
                prefix = _get_numbering_prefix(para, numbering_cache, numbering_state)
                if prefix and not paragraph_text.startswith(prefix):
                    paragraph_text = f"{prefix} {paragraph_text}".strip()
            pid = get_paragraph_id(para, fallback_idx=i)
            items.append({
                'text': paragraph_text.strip(),
                'para_id': pid,
                'para_index0': i - 1  # zero-based index into doc.paragraphs
            })

    return items

def split_docx_to_question_with_ids(file_path: str, second_split: bool = False, not_nbu: bool = False) -> List[Dict]:
    dprint("split_docx_to_question_with_ids()")

    """
    Split into numbered questions and keep a list of paragraph IDs per question.
    Returns a list of dicts: {'question_text': str, 'para_ids': [str, ...], 'para_indices0': [int, ...]}.
    """

    items = extract_text_with_font_and_ids(file_path, not_nbu)

    if len(items) == 0:
        print(f"len(items): {len(items)} --- FIRST RESULT NOT GIVE QUESTION, MAKING SECOND TRY WITHOUT SEARCH FOR 'Загальні положення")
        items = extract_text_with_font_and_ids(file_path, True)


    # Top-level markers like "117." or "308-1."
    top_marker = re.compile(r'^\d+(?:-\d+)?\.\s')

    # Build top-level buckets from consecutive paragraphs
    buckets_parts = []  # each item: {'parts': [{'text', 'pid', 'idx0'}, ...]}
    current = {'parts': []}
    for it in items:
        txt = it['text']
        if top_marker.match(txt):
            if current['parts']:
                buckets_parts.append(current)
            current = {'parts': []}
        current['parts'].append({'text': txt, 'pid': it['para_id'], 'idx0': it.get('para_index0')})
    if current['parts']:
        buckets_parts.append(current)

    # Helper to turn parts -> bucket object
    def build_bucket(parts):
        question_text = normalize_multiline_enumeration('\n'.join(p['text'] for p in parts))
        return {
            'question_text': question_text,
            'para_ids': [p['pid'] for p in parts],
            'para_indices0': [p['idx0'] for p in parts],
            '_parts': parts,   # keep for optional second split
        }

    buckets = [build_bucket(b['parts']) for b in buckets_parts]

    if not second_split:
        # Strip helper key before returning
        for b in buckets:
            b.pop('_parts', None)
        return buckets

    # Second-level markers like "1.1.", "2.3.4." (at least one dot level)
    sub_marker_block = re.compile(
        r'(?s)((?:\d+(?:\.\d+)+\.)\s.*?)(?=(?:\n\d+(?:\.\d+)+\.\s)|\Z)'
    )

    refined = []
    for b in buckets:
        parts = b['_parts']
        # Build joined text and paragraph spans (start,end) for overlap calc
        joined = ''
        spans = []  # list of (start, end, pid, idx0)
        pos = 0
        for p in parts:
            t = p['text']
            start = pos
            joined += t
            pos += len(t)
            end = pos
            spans.append((start, end, p['pid'], p['idx0']))
            # add the newline we used when joining for detection between paragraphs
            joined += '\n'
            pos += 1

        matches = list(re.finditer(sub_marker_block, joined))

        if not matches:
            # No sub-splits detected — keep the whole bucket
            refined.append({
                'question_text': b['question_text'],
                'para_ids': b['para_ids'],
                'para_indices0': b['para_indices0'],
            })
            continue

        # Include the top-level question itself (paragraphs before first sub-marker)
        top_level_parts = []
        for p in parts:
            if sub_marker_block.match(p['text']):
                break
            top_level_parts.append(p)
        if top_level_parts:
            refined.append({
                'question_text': normalize_multiline_enumeration('\n'.join(pt['text'] for pt in top_level_parts)),
                'para_ids': [pt['pid'] for pt in top_level_parts],
                'para_indices0': [pt['idx0'] for pt in top_level_parts],
            })

        # Build refined sub-questions with their OWN para_ids by overlap
        for m in matches:
            mstart, mend = m.span()
            sub_ids, sub_idxs = [], []
            for s, e, pid, idx in spans:
                # overlap if paragraph span intersects match span
                if s < mend and e > mstart:
                    sub_ids.append(pid)
                    sub_idxs.append(idx)
            refined.append({
                'question_text': normalize_multiline_enumeration(m.group(1)),
                'para_ids': sub_ids,
                'para_indices0': sub_idxs,
            })

    return refined


# Print runs for given paragraph IDs; if IDs are missing in the document, fall back to indices.
# para_indices0 are zero-based indices into doc.paragraphs.
def print_runs_for_paras(docx_path, para_ids=None, para_indices0=None):
    dprint("print_runs_for_paras()")

    """Print runs for given paragraph IDs; if IDs are missing in the document, fall back to indices.
    para_indices0 are zero-based indices into doc.paragraphs.
    """
    doc = Document(docx_path)

    # Fast path: if indices provided, use them directly
    if para_indices0:
        for idx in para_indices0:
            if 0 <= idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                pid = para._element.get(qn('w14:paraId')) or para._element.get(qn('w:paraId'))
                # if DEBUG:
                #     print(f"--- Paragraph idx0={idx} pid={pid} ---")
                #     print(para.text)

        return

    # Build map from IDs to paragraph
    if para_ids:
        id_to_para = {}
        for para in doc.paragraphs:
            pid = para._element.get(qn('w14:paraId')) or para._element.get(qn('w:paraId'))
            if pid:
                id_to_para[pid] = para
        for pid in para_ids:
            para = id_to_para.get(pid)
            if not para:
                continue
            # if DEBUG:
            #     print(f"--- Paragraph ID: {pid} ---")
            #     print(para.text)


if __name__ == "__main__":
    print(f"{'='*33} if __name__ == __main__: read_split_nbu.py {'='*33}")

    # if sys.gettrace():  # True if running under debugger
    # sys.argv = ["read_split_nbu.py", "DORA_req.docx", "False"]

    if len(sys.argv) > 1:
        file_name = sys.argv[1]

    if len(sys.argv) > 2:
        not_nbu = sys.argv[2].strip().lower() in ("true", "1", "yes", "y")
    else:
        not_nbu = False

    

    nbu_folder_path = 'NBU/'


    # file_name = 'NBU64.docx'

    file_path = os.path.join(nbu_folder_path, file_name)

    # questions = split_docx_to_question(file_path, True)
    questions_struct = split_docx_to_question_with_ids(file_path, True, not_nbu)

    if len(questions_struct) == 0:
        print("FIRST RESULT NOT GIVE QUESTION, MAKING SECOND TRY WITHOUT SEARCH FOR 'Загальні положення")
        questions_struct = split_docx_to_question_with_ids(file_path, True, True)


    # doc = Document(file_path)



    # for i, q in enumerate(questions_struct):
    #     print(f"q['para_ids'] : {q['para_ids']}")
    #     print(f"q['para_indices0'] : {q.get('para_indices0')}")
    #     # Prefer IDs if the doc has them; otherwise fall back to indices
    #     print_runs_for_paras(file_path, para_ids=q['para_ids'], para_indices0=q.get('para_indices0'))
    #     # print(q)
    #     # print(q['para_ids'])
    #     # for id in q['para_ids']:
    #     #     print(f"paragraph id : {id}")
    #     #     print(doc.paragraphs[id])
    #     #     print("-"*50)
    #     if i > 10:
    #         break



    # Write plain text file (back-compat)
    with open(f"questions_new_gpt_{file_name[:-5]}.txt", "w", encoding='utf-8') as f:
        for q in questions_struct:
            f.write(q['question_text'])
            f.write("\n")
            f.write("-"*120)
            f.write("\n")



    # Write JSON with paragraph IDs for precise mapping back to the docx
    with open(f"questions_new_gpt_{file_name[:-5]}_with_ids.json", "w", encoding='utf-8') as jf:
        json.dump(questions_struct, jf, ensure_ascii=False, indent=2)

    print(f"\nlen of questions : {len(questions_struct)}")
    print(sys.argv)
    print(f"not_nbu : {not_nbu}")

