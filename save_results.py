from datetime import datetime
from typing import List, Tuple
import uuid
import os
import re

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
# CommentsPart import path differs across python-docx builds; try both
try:
    from docx.parts.comments import CommentsPart
except Exception:
    try:
        from docx.parts.comments_part import CommentsPart  # type: ignore
    except Exception:
        CommentsPart = None  # type: ignore
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Optional debug gating via Config.DEBUG
try:
    from config import Config
    DEBUG = getattr(Config, "DEBUG", False)
except Exception:
    DEBUG = False

def dprint(msg: str):
    if DEBUG:
        print(msg)

# Helper for writing multi-line, lightly-marked-up comment text into comments.xml
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

def sanitize_markdown(text: str) -> str:
    """Normalize markdown-ish text from the LLM for DOCX comments."""
    dprint("sanitize_markdown()")
    if text is None:
        return ""
    # strip NULs that sometimes appear after emojis like '✅\x00'
    text = text.replace("\x00", "")
    # normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # collapse excessive trailing spaces
    lines = [ln.rstrip() for ln in text.split("\n")]
    return "\n".join(lines)

def _append_markdownish_comment_text(container: OxmlElement, text: str) -> None:
    dprint("_append_markdownish_comment_text()")
    """Append LLM text to a comments.xml container as multiple <w:p> elements.
    Supports a *light* subset of Markdown:
      - Headings (#, ##, ###) rendered as bold lines
      - Bulleted items starting with '- ' or '* ' rendered with a '• ' prefix
      - **bold** and *italic* inline markup
    Newlines produce new paragraphs. Empty lines create empty paragraphs.
    """
    import re

    text = sanitize_markdown(text or "")
    lines = text.split("\n")

    bold_pat = re.compile(r"\*\*(.+?)\*\*")
    ital_pat = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
    heading_pat = re.compile(r"^\s{0,3}(#{1,6})\s+(.*)$")
    bullet_pat  = re.compile(r"^\s{0,3}[-\*]\s+(.*)$")

    def _add_run(p: OxmlElement, s: str, bold: bool = False, italic: bool = False):
        if s == "":
            return
        r = OxmlElement("w:r")
        if bold or italic:
            rpr = OxmlElement("w:rPr")
            if bold:
                rpr.append(OxmlElement("w:b"))
            if italic:
                rpr.append(OxmlElement("w:i"))
            r.append(rpr)
        t = OxmlElement("w:t")
        t.set(XML_SPACE, "preserve")
        t.text = s
        r.append(t)
        p.append(r)

    def _emit_inline_tokens(p: OxmlElement, content: str, force_bold: bool = False):
        i = 0
        while i < len(content):
            m_b = bold_pat.search(content, i)
            m_i = ital_pat.search(content, i)
            m = None
            style = "t"
            if m_b and (not m_i or m_b.start() < m_i.start()):
                m, style = m_b, "b"
            elif m_i:
                m, style = m_i, "i"
            if m:
                if m.start() > i:
                    _add_run(p, content[i:m.start()], bold=force_bold)
                if style == "b":
                    _add_run(p, m.group(1), bold=True or force_bold)
                else:
                    _add_run(p, m.group(1), bold=force_bold, italic=True)
                i = m.end()
            else:
                _add_run(p, content[i:], bold=force_bold)
                break

    for ln in lines:
        p = OxmlElement("w:p")
        if not ln.strip():
            # empty paragraph
            container.append(p)
            continue

        # Headings (# ...), render the text portion in bold
        m_h = heading_pat.match(ln)
        if m_h:
            level, content = m_h.group(1), m_h.group(2)
            _emit_inline_tokens(p, content, force_bold=True)
            container.append(p)
            continue

        # Bullets (- ... or * ...), add a '• ' prefix
        m_bul = bullet_pat.match(ln)
        if m_bul:
            _add_run(p, "• ")
            _emit_inline_tokens(p, m_bul.group(1))
            container.append(p)
            continue

        # Plain paragraph
        _emit_inline_tokens(p, ln)
        container.append(p)

def add_comment_to_paragraphs(doc, para_ids, comment_text, author="AI helper", initials="AI", color=None):
    dprint(f"add_comment_to_paragraphs(len={len(para_ids) if para_ids else 0}, color={color})")
    """
    Insert a DOCX comment spanning the paragraph indices in para_ids (zero-based).

    Note: despite the historical name, this function expects a list of paragraph
    indices, not Word paragraph IDs. It will color all runs in the range if a
    color or verdict emoji is provided.
    """
    if not para_ids:
        return

    paragraphs = doc.paragraphs

    comment_id = str(uuid.uuid4().int)[:8]

    # Set color if needed
    if color:
        color_map = {
            "❌": RGBColor(206, 43, 57),
            "✅": RGBColor(94, 140, 97),
            "❓": RGBColor(194, 146, 25),
            "red": RGBColor(206, 43, 57),
            "green": RGBColor(94, 140, 97),
            "yellow": RGBColor(194, 146, 25),
        }
        key = color.lower() if isinstance(color, str) else color
        color_rgb = color_map.get(key, RGBColor(0, 0, 0))

        for idx in para_ids:
            for run in doc.paragraphs[idx].runs:
                run.font.color.rgb = color_rgb

    # Ensure start/end paragraphs have at least one run; python-docx cannot
    # insert comment markers against empty paragraphs.
    def _ensure_run(par):
        if not par.runs:
            par.add_run("")
        return par.runs

    start_par = doc.paragraphs[para_ids[0]]
    end_par = doc.paragraphs[para_ids[-1]]
    _ensure_run(start_par)
    _ensure_run(end_par)

    start_run = start_par.runs[0]._r
    end_run = end_par.runs[-1]._r

    comment_start = OxmlElement("w:commentRangeStart")
    comment_start.set(qn("w:id"), comment_id)
    start_run.addprevious(comment_start)

    comment_end = OxmlElement("w:commentRangeEnd")
    comment_end.set(qn("w:id"), comment_id)
    end_run.addnext(comment_end)

    comment_ref = OxmlElement("w:commentReference")
    comment_ref.set(qn("w:id"), comment_id)
    end_run.addnext(comment_ref)

    # Create comment part if needed
    part = doc.paragraphs[para_ids[0]].part
    try:
        comments_part = part._comments_part
    except AttributeError:
        if CommentsPart is None:
            # Fallback: if comments API is unavailable, skip creating comment part
            # Color was already applied above, so at least the text is marked.
            dprint("CommentsPart unavailable in this python-docx build; skipping comment creation")
            return
        comments_part = CommentsPart(part.package, PackURI("/word/comments.xml"), RT.COMMENTS)
        part._comments_part = comments_part
        part.relate_to(comments_part, RT.COMMENTS)

    comments = comments_part._element

    # Use ISO 8601 UTC timestamp
    formatted_time = datetime.utcnow().isoformat() + "Z"

    new_comment = OxmlElement("w:comment")
    new_comment.set(qn("w:author"), author)
    new_comment.set(qn("w:initials"), initials)
    new_comment.set(qn("w:date"), formatted_time)
    new_comment.set(qn("w:id"), comment_id)

    # Add comment content with preserved formatting/newlines
    _append_markdownish_comment_text(new_comment, comment_text)
    comments.append(new_comment)



def add_comment_to_text_range(doc, target_text, comment_text, author="AI helper", initials="AI", color=None):
    dprint("add_comment_to_text_range()")
    full_text = ""
    para_map = []

    # Build full document text + map paragraphs
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            full_text += para.text + "\n"
            para_map.append((idx, para.text))

    # Find which paragraphs contain the target_text
    buffer = ""
    matched_indices = []
    for idx, text in para_map:
        if text in target_text:
            buffer += text + "\n"
            matched_indices.append(idx)
        if target_text in buffer:
            # print(f"buffer: {buffer}")
            break

    if not matched_indices:
        dprint("No paragraph range matched target_text")
        return

    # Use first and last paragraph in the matched range
    start_idx = matched_indices[0]
    end_idx = matched_indices[-1]
    comment_id = str(uuid.uuid4().int)[:8]

    # Set color if needed
    if color:
        color_map = {
            "❌": RGBColor(206, 43, 57),
            "✅": RGBColor(94, 140, 97),
            "❓": RGBColor(194, 146, 25),
            "red": RGBColor(206, 43, 57),
            "green": RGBColor(94, 140, 97),
            "yellow": RGBColor(194, 146, 25),
        }
        key = color.lower() if isinstance(color, str) else color
        color_rgb = color_map.get(key, RGBColor(0, 0, 0))

        for i in range(start_idx, end_idx + 1):
            for run in doc.paragraphs[i].runs:
                run.font.color.rgb = color_rgb

    start_run = None
    end_run = None

    start_run = doc.paragraphs[start_idx].runs[0]._r
    end_run = doc.paragraphs[end_idx].runs[-1]._r

    comment_start = OxmlElement("w:commentRangeStart")
    comment_start.set(qn("w:id"), comment_id)
    start_run.addprevious(comment_start)

    comment_end = OxmlElement("w:commentRangeEnd")
    comment_end.set(qn("w:id"), comment_id)
    end_run.addnext(comment_end)

    comment_ref = OxmlElement("w:commentReference")
    comment_ref.set(qn("w:id"), comment_id)
    end_run.addnext(comment_ref)

    # Create comment part if needed
    part = doc.paragraphs[start_idx].part
    try:
        comments_part = part._comments_part
    except AttributeError:
        if CommentsPart is None:
            dprint("CommentsPart unavailable in this python-docx build; skipping comment creation")
            return
        comments_part = CommentsPart(part.package, PackURI("/word/comments.xml"), RT.COMMENTS)
        part._comments_part = comments_part
        part.relate_to(comments_part, RT.COMMENTS)

    comments = comments_part._element

    new_comment = OxmlElement("w:comment")
    new_comment.set(qn("w:author"), author)
    new_comment.set(qn("w:initials"), initials)
    new_comment.set(qn("w:date"), datetime.utcnow().isoformat() + "Z")
    new_comment.set(qn("w:id"), comment_id)

    # Add comment content with preserved formatting/newlines
    _append_markdownish_comment_text(new_comment, comment_text)
    comments.append(new_comment)

def add_comment(paragraph, text_to_comment, comment_text, author="Author", initials="AU", color=None):
    dprint("add_comment()")
    # Search for text
    if text_to_comment not in paragraph.text:
        return

    # Split text around the target
    parts = paragraph.text.split(text_to_comment)
    paragraph.clear()

    # Add first part
    paragraph.add_run(parts[0])

    # Add the text with comment marker
    run = paragraph.add_run(text_to_comment)

    # Set color if specified
    if color:
        color_map = {
            "red": RGBColor(206, 43, 57),
            "green": RGBColor(94, 140, 97),
            "yellow": RGBColor(194, 146, 25),
        }
        run.font.color.rgb = color_map.get(color.lower(), RGBColor(0, 0, 0)) # black is default

    comment_id = str(uuid.uuid4().int)[:8]  # unique ID for comment

    # Create comment range start
    comment_start = OxmlElement('w:commentRangeStart')
    comment_start.set(qn('w:id'), comment_id)
    run._r.addprevious(comment_start)

    # Create comment range end
    comment_end = OxmlElement('w:commentRangeEnd')
    comment_end.set(qn('w:id'), comment_id)
    run._r.addnext(comment_end)

    # Add comment reference
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), comment_id)
    run._r.addnext(comment_ref)

    # Add remaining text
    paragraph.add_run(parts[1])

    # Now add the comment itself to the comments part
    part = paragraph.part
    try:
        comments_part = part._comments_part
    except AttributeError:
        # Create comments part if it doesn't exist
        comments_part = CommentsPart(part.package, PackURI("/word/comments.xml"), RT.COMMENTS)
        part._comments_part = comments_part
        part.relate_to(comments_part, RT.COMMENTS)

    comments = comments_part._element

    new_comment = OxmlElement('w:comment')
    new_comment.set(qn('w:author'), author)
    new_comment.set(qn('w:initials'), initials)
    new_comment.set(qn('w:date'), datetime.utcnow().isoformat() + "Z")
    new_comment.set(qn('w:id'), comment_id)

    # Add comment content with preserved formatting/newlines
    _append_markdownish_comment_text(new_comment, comment_text)
    comments.append(new_comment)


def read_docx_by_paragraph(file_path):
    dprint("read_docx_by_paragraph()")
    if not os.path.exists(file_path):
        raise FileNotFoundError("File with this path does not exists")

    doc = Document(file_path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            yield para
